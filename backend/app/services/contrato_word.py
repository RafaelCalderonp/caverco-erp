import base64
import io
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MESES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _fecha_larga(d: date | None) -> str:
    if not d:
        return ""
    return f"{d.day:02d} {MESES[d.month - 1]} {d.year}"


def _clp(valor) -> str:
    if valor is None:
        return ""
    return f"${int(valor):,}".replace(",", ".") + ".-"


def _parrafo(doc, partes, bold_default=False, align=None, space_after=10):
    """partes: lista de (texto, bold) o strings sueltos (heredan bold_default)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    for parte in partes:
        if isinstance(parte, tuple):
            texto, bold = parte
        else:
            texto, bold = parte, bold_default
        run = p.add_run(texto)
        run.bold = bold
        run.font.size = Pt(11)
    return p


def _logo_bytes(logo_url: str | None) -> bytes | None:
    if not logo_url or not logo_url.startswith("data:"):
        return None
    try:
        b64 = logo_url.split(",", 1)[1]
        return base64.b64decode(b64)
    except Exception:
        return None


def generar_contrato_docx(empresa, empleado, contrato, cargo_nombre, obra, afp_nombre, isapre_nombre, tipo_contrato_codigo) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))

    _parrafo(doc, ["CONTRATO DE TRABAJO"], bold_default=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    nombre_completo = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()

    _parrafo(doc, [
        f"En {empresa.ciudad or 'Santiago'}, a ", (_fecha_larga(contrato.fecha_contrato), True),
        ", entre ", (empresa.razon_social, True), ", RUT N° ", (empresa.rut, True),
        ", representado por ", (empresa.representante_legal or "", True),
        ", RUT ", (empresa.rut_representante_legal or "", True),
        f", domiciliado en {empresa.direccion or ''} comuna de {empresa.comuna or ''}, ciudad de {empresa.ciudad or ''}, correo electrónico ",
        (empresa.email or "", True),
        ' que en adelante se denominará "EL EMPLEADOR", y don/doña ', (nombre_completo, True),
        ", cédula nacional de identidad N° ", (empleado.rut, True),
        ", nacionalidad ", (empleado.nacionalidad or "Chilena", True),
        ", nacido el ", (_fecha_larga(empleado.fecha_nacimiento), True),
        ", estado civil ", (empleado.estado_civil or "", True),
        ", domiciliado en ", (empleado.direccion or "", True), ", comuna de ", (empleado.comuna or "", True),
        ", región ", (empleado.region or "", True),
        ", correo electrónico ", (empleado.email_personal or empleado.email_corporativo or "", True),
        ' en adelante "EL TRABAJADOR" se ha convenido el siguiente contrato de trabajo, en adelante el "CONTRATO":',
    ])

    _parrafo(doc, [
        ("PRIMERO: ", True),
        "El empleador contrata al Trabajador para que preste los servicios de ", (cargo_nombre or "", True),
        ", su tarea principal es la ejecución de las labores asociadas a dicho cargo, sea esta principal, "
        "complementaria o alternativa. No cumplir y no ajustarse a la pauta implica un incumplimiento grave "
        "de las obligaciones que impone el contrato.",
    ])

    if obra:
        lugar = f"{obra.direccion or obra.nombre}, comuna de {obra.comuna or ''}, región {obra.region or ''}"
    else:
        lugar = f"{empresa.direccion or ''}, comuna de {empresa.comuna or ''}"
    _parrafo(doc, [
        ("SEGUNDO: ", True),
        "Los servicios serán prestados en el establecimiento ubicado en ", (lugar, True), ".",
    ])

    if contrato.horario_detalle:
        _parrafo(doc, [
            ("TERCERO: ", True),
            f"La jornada de trabajo será de ", (str(contrato.horas_semanales or 42), True),
            f" horas semanales, bajo jornada ", (contrato.jornada or "Completa", True),
            ", distribuidas de la siguiente manera: ", (contrato.horario_detalle, True), ".",
        ])
    else:
        _parrafo(doc, [
            ("TERCERO: ", True),
            f"La jornada de trabajo será de ", (str(contrato.horas_semanales or 42), True),
            f" horas semanales, distribuidas conforme a las necesidades de la empresa, bajo jornada ",
            (contrato.jornada or "Completa", True), ".",
        ])

    _parrafo(doc, [
        ("CUARTO: ", True),
        "El EMPLEADOR se compromete a remunerar los servicios del TRABAJADOR con un sueldo base mensual de ",
        (_clp(contrato.sueldo_bruto), True), " ",
    ])
    _parrafo(doc, [
        "Además, la Empresa pagará al trabajador una gratificación equivalente a un 25% de las remuneraciones. "
        "La gratificación indicada se pagará mensualmente en sustitución a la señalada en el Art. 47 del "
        "Código del Trabajo, según establece el Art. 50 del cuerpo legal.",
    ])
    _parrafo(doc, [
        "Las remuneraciones se pagarán, dentro de la hora siguiente del término de la jornada laboral, en "
        "moneda nacional, el día cinco hábil del mes siguiente al que se devengaron las remuneraciones "
        "respectivas, en dinero efectivo o, por solicitarlo el trabajador en este acto, por medio de una "
        "chequera electrónica, vale vista o transferencia bancaria. Del monto de las remuneraciones, la "
        "Empresa hará las deducciones que establecen las leyes vigentes o que sean autorizadas por EL "
        "TRABAJADOR, por escrito. Todo sin perjuicio de los descuentos propios por concepto de tiempo no "
        "trabajado, debido a inasistencias, permisos y atrasos.",
    ])

    colacion = int(contrato.colacion or 0)
    movilizacion = int(contrato.movilizacion or 0)
    if colacion > 0 or movilizacion > 0:
        beneficios = []
        if colacion > 0:
            beneficios.append(f"a) Asignación de colación: {_clp(colacion)} mensuales")
        if movilizacion > 0:
            beneficios.append(f"{'b' if colacion > 0 else 'a'}) Asignación de movilización: {_clp(movilizacion)} mensuales")
        _parrafo(doc, [
            ("CUARTO BIS: ", True),
            "Adicionalmente a la remuneración pactada, el EMPLEADOR otorgará al TRABAJADOR los siguientes "
            "beneficios de carácter NO REMUNERACIONAL conforme al artículo 41 inciso 2° del Código del Trabajo: ",
            "; ".join(beneficios) + ". ",
            "Dichos montos no constituyen remuneración para efectos de cotizaciones previsionales ni de salud.",
        ])

    if tipo_contrato_codigo == "PLAZO_FIJO":
        _parrafo(doc, [
            ("QUINTO: ", True),
            "El presente contrato tiene el carácter de Plazo Fijo y durará hasta el ",
            (_fecha_larga(contrato.fecha_termino_pactada), True),
            ". Las partes pueden ponerle término, además de común acuerdo, en la forma, las condiciones y "
            "las causales que señalan los artículos 159, 160 y 161 del Código del Trabajo.",
        ])
    else:
        _parrafo(doc, [
            ("QUINTO: ", True),
            "El presente contrato tiene el carácter de Indefinido. Las partes pueden ponerle término, "
            "además de común acuerdo; y una de ellas, en la forma, las condiciones y las causales que "
            "señalan los artículos 159, 160 y 161 del Código del Trabajo.",
        ])

    _parrafo(doc, [
        ("SEXTO: ", True),
        "Será obligatorio por parte del trabajador el uso de todos los elementos y equipos de protección "
        "personal que le proporcione la empresa. El no cumplimiento de esta obligación será considerado un "
        "incumplimiento grave de las obligaciones contractuales, además de exponer la vida y la salud del "
        "trabajador. La empresa se ve obligada por la normativa vigente, a sancionar al trabajador que "
        "incumpla esta obligación.",
    ])

    _parrafo(doc, [
        ("SÉPTIMO: ", True),
        "El Trabajador declara que su régimen previsional corresponde a la AFP ", (afp_nombre or "", True),
        ", mientras que para efectos de su cotización de salud se encuentra en ", (isapre_nombre or "", True), ".",
    ])

    _parrafo(doc, [
        ("OCTAVO: ", True),
        "El trabajador declara que ingresó al servicio del empleador a contar del día ",
        (_fecha_larga(contrato.fecha_inicio), True), ".",
    ])

    _parrafo(doc, [
        ("NOVENO: ", True),
        "El presente CONTRATO se firma en tres ejemplares del mismo tenor, fecha y valor probatorio, "
        "declarando EL TRABAJADOR haber recibido un ejemplar de él, y que éste es fiel reflejo de la "
        "relación laboral existente entre las partes.",
    ])

    _parrafo(doc, ["En comprobante, y previa lectura firman."], space_after=40)

    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = True
    celda_empresa, celda_trabajador = tabla.rows[0].cells
    celda_empresa.text = f"{empresa.razon_social}\nRUT: {empresa.rut}"
    celda_trabajador.text = f"{nombre_completo}\nRUT: {empleado.rut}"
    for celda in (celda_empresa, celda_trabajador):
        for p in celda.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generar_anexo_docx(empresa, empleado, contrato, anexo, tipo_anexo_codigo, cargo_nombre) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))

    _parrafo(doc, ["ANEXO DE CONTRATO"], bold_default=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    nombre_completo = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()

    _parrafo(doc, [
        f"En {empresa.ciudad or 'Santiago'}, a ", (_fecha_larga(anexo.fecha_anexo), True),
        ", entre ", (empresa.razon_social, True), ", RUT N° ", (empresa.rut, True),
        f", domiciliado en {empresa.direccion or ''} comuna de {empresa.comuna or ''}, ciudad de {empresa.ciudad or ''}, correo electrónico ",
        (empresa.email or "", True),
        ' que en adelante se denominará "EL EMPLEADOR", y don/doña ', (nombre_completo, True),
        ", cédula nacional de identidad N° ", (empleado.rut, True),
        ", nacionalidad ", (empleado.nacionalidad or "Chilena", True),
        ", nacido el ", (_fecha_larga(empleado.fecha_nacimiento), True),
        ", de estado civil ", (empleado.estado_civil or "", True),
        ", domiciliado en ", (empleado.direccion or "", True), ", comuna de ", (empleado.comuna or "", True),
        ", región ", (empleado.region or "", True),
        ", correo electrónico ", (empleado.email_personal or empleado.email_corporativo or "", True),
        ' en adelante "EL TRABAJADOR" se ha convenido el siguiente contrato de trabajo, en adelante el "ANEXO DE CONTRATO":',
    ])

    _parrafo(doc, ["COMPARECENCIA:"], bold_default=True, space_after=6)
    _parrafo(doc, [
        "Con Fecha ", (_fecha_larga(contrato.fecha_contrato), True),
        " Las partes declaran que han suscrito un contrato de trabajo en virtud del cual don/doña ",
        (nombre_completo, True), ", ha prestado servicios en calidad de ", (cargo_nombre or "", True),
        ', para la empresa, en adelante el "Contrato de trabajo" y que se encuentra plenamente vigente.',
    ])
    _parrafo(doc, [
        "Además, Las partes acuerdan de esta fecha que se modificará la cláusula Segunda del contrato original en lo siguiente.",
    ])

    _parrafo(doc, ["CLÁUSULAS:"], bold_default=True, space_after=6)

    if tipo_anexo_codigo == "PRORROGA_PLAZO":
        fecha_anterior = (anexo.valor_anterior or {}).get("fecha_termino_pactada")
        fecha_nueva = (anexo.valor_nuevo or {}).get("fecha_termino_pactada")
        fecha_anterior_dt = date.fromisoformat(fecha_anterior) if fecha_anterior else None
        fecha_nueva_dt = date.fromisoformat(fecha_nueva) if fecha_nueva else None
        if fecha_anterior_dt is None and anexo.fecha_anexo:
            from datetime import timedelta
            fecha_anterior_dt = anexo.fecha_anexo - timedelta(days=1)
        if fecha_nueva_dt is None:
            fecha_nueva_dt = contrato.fecha_termino_pactada
        _parrafo(doc, [("PRIMERO: Modificación de duración Contrato: ", True)])
        _parrafo(doc, [
            'Cláusula que establece: "Las partes acuerdan prorrogar el contrato de trabajo que vence el ',
            (_fecha_larga(fecha_anterior_dt), True),
            ', hasta el ', (_fecha_larga(fecha_nueva_dt), True), '.',
        ])
    elif tipo_anexo_codigo == "CONV_INDEFINIDO":
        _parrafo(doc, [("PRIMERO: Modificación de duración Contrato: ", True)])
        _parrafo(doc, [
            'Cláusula que establece: "Las partes vienen a acordar que el contrato de trabajo suscrito con fecha ',
            (_fecha_larga(contrato.fecha_contrato), True),
            ", originalmente pactado a plazo fijo, se modifica en cuanto a su duración, pasando a tener el "
            "carácter de contrato de trabajo de duración indefinida a contar de la fecha de suscripción del "
            "presente Anexo. En todo lo no modificado por el presente instrumento, se mantienen plenamente "
            "vigentes las demás cláusulas y condiciones del contrato original.",
        ])
    else:
        _parrafo(doc, [("PRIMERO: ", True), (anexo.observacion or "", True)])

    _parrafo(doc, [("SEGUNDO: Final", True)])
    _parrafo(doc, ["En todo lo no modificado por el presente instrumento, seguirá rigiendo el Contrato de Trabajo."])
    _parrafo(doc, ["El presente instrumento se firma en dos ejemplares del mismo tenor, quedando uno en poder de cada parte."], space_after=40)

    tabla = doc.add_table(rows=1, cols=2)
    tabla.autofit = True
    celda_empresa, celda_trabajador = tabla.rows[0].cells
    celda_empresa.text = f"{empresa.razon_social}\nRUT: {empresa.rut}"
    celda_trabajador.text = f"{nombre_completo}\nRUT: {empleado.rut}"
    for celda in (celda_empresa, celda_trabajador):
        for p in celda.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─── ENTREGA DE REGLAMENTO INTERNO ───────────────────────────────────────────

MESES_RRHH = ["enero","febrero","marzo","abril","mayo","junio",
               "julio","agosto","septiembre","octubre","noviembre","diciembre"]


def generar_reglamento_docx(empresa, empleado, fecha_entrega: date) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.59)
    section.page_height   = Cm(27.94)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    nombre_empleado = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()
    cargo_nombre    = getattr(empleado, "cargo_nombre", "") or ""
    empresa_nombre  = empresa.razon_social or ""

    # Logo
    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))

    # Título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(16)
    r = p_titulo.add_run(
        "FORMULARIO RECEPCIÓN DEL REGLAMENTO INTERNO DE ORDEN, HIGIENE Y SEGURIDAD, "
        "ART. 67º DE LA LEY Nº 16.744, TITULO III DEL CÓDIGO DEL TRABAJO, D.F.L. N° 1"
    )
    r.bold = True
    r.font.size = Pt(11)

    # Párrafos
    for txt in [
        f"Declaro haber recibido en forma gratuita una copia del reglamento interno de orden, higiene y seguridad "
        f"de la empresa {empresa_nombre} de acuerdo a lo establecido en el Art. 56 del D.S N° 40 y Art. 67 de la Ley 16.744.",
        "Asumo mi responsabilidad de dar lectura a su contenido y dar cumplimiento a las obligaciones, prohibiciones, "
        "normas de orden, higiene y seguridad que en él están escritas, como así también a las disposiciones "
        "establecidas por el Organismo Administrador del Seguro de Accidentes del Trabajo y Enfermedades "
        "Profesionales a que se encuentre adherida la empresa.",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(txt)
        r.font.size = Pt(10)

    doc.add_paragraph("- " * 55).runs[0].font.size = Pt(8)

    # Tabla datos
    fecha_str = f"{fecha_entrega.day:02d} {MESES_RRHH[fecha_entrega.month-1]} {fecha_entrega.year}" if fecha_entrega else ""
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    for i, (lbl, val) in enumerate([
        ("Nombre completo",       nombre_empleado),
        ("R.U.T.",                empleado.rut or ""),
        ("Sección",               cargo_nombre),
        ("Firma del trabajador",  ""),
        ("Fecha de entrega",      fecha_str),
    ]):
        tbl.rows[i].cells[0].text = lbl
        tbl.rows[i].cells[1].text = val
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True if tbl.rows[i].cells[0].paragraphs[0].runs else False
        tbl.rows[i].cells[0].width = Cm(5.5)
        tbl.rows[i].cells[1].width = Cm(11.5)
        if lbl == "Firma del trabajador":
            tbl.rows[i].height = Cm(2.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─── ENTREGA DE EPP ───────────────────────────────────────────────────────────

def _set_cell_bg_epp(cell, hex_color: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text="", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if text:
        r = p.add_run(str(text))
        r.bold = bold
        r.font.size = Pt(size)


def generar_epp_docx(empresa, empleado, entrega) -> bytes:
    from docx.enum.table import WD_ALIGN_VERTICAL
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.59)
    section.page_height   = Cm(27.94)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)

    nombre_empleado = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()
    fecha_entrega   = entrega.fecha_entrega
    entregado_por   = entrega.entregado_por or "Salvador Calderón"
    items           = entrega.items or []

    # ── Encabezado: logo + título + folio ─────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=3)
    hdr.style = "Table Grid"

    logo_cell = hdr.rows[0].cells[0]
    logo_cell.width = Cm(3.5)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_cell.text = ""
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p_logo.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))
    else:
        r = p_logo.add_run(getattr(empresa, "razon_social", "") or "")
        r.bold = True
        r.font.size = Pt(10)

    title_cell = hdr.rows[0].cells[1]
    title_cell.width = Cm(13.0)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_text(title_cell, "REGISTRO DE ENTREGA DE ELEMENTOS DE PROTECCIÓN PERSONAL",
               bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    folio_cell = hdr.rows[0].cells[2]
    folio_cell.width = Cm(3.0)
    _cell_text(folio_cell, f"Folio N°\n{entrega.folio or ''}", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Texto legal ───────────────────────────────────────────────────────────
    p_ley = doc.add_paragraph()
    p_ley.paragraph_format.space_after = Pt(8)
    r1 = p_ley.add_run('De acuerdo a lo estipulado en la ')
    r1.font.size = Pt(9)
    r2 = p_ley.add_run('Ley 16.744 Art. 68 inciso tres')
    r2.bold = True; r2.font.size = Pt(9)
    r3 = p_ley.add_run(
        ' "Las empresas deberán proporcionar a sus trabajadores los equipos e implementos de '
        'protección necesarios no pudiendo en caso alguno cobrarles su valor".'
    )
    r3.font.size = Pt(9); r3.italic = True

    # ── Datos del trabajador ──────────────────────────────────────────────────
    datos = doc.add_table(rows=4, cols=2)
    datos.style = "Table Grid"
    filas_datos = [
        ("Nombre del trabajador", nombre_empleado),
        ("Cédula de Identidad",   empleado.rut or ""),
        ("Cargo",                 getattr(empleado, "cargo_nombre", "") or ""),
        ("Área / Obra / Faena",   empresa.razon_social or ""),
    ]
    for i, (lbl, val) in enumerate(filas_datos):
        _cell_text(datos.rows[i].cells[0], lbl, bold=True, size=9)
        _cell_text(datos.rows[i].cells[1], val, size=9)
        datos.rows[i].cells[0].width = Cm(5.0)
        datos.rows[i].cells[1].width = Cm(14.5)

    doc.add_paragraph()

    # ── Subtítulo ─────────────────────────────────────────────────────────────
    p_sub = doc.add_paragraph("REGISTRO PERSONAL DE ENTREGA DE EPP")
    p_sub.runs[0].bold = True
    p_sub.runs[0].font.size = Pt(10)
    p_sub.paragraph_format.space_after = Pt(4)

    # ── Tabla EPP — 6 columnas ────────────────────────────────────────────────
    n_rows = max(len(items), 8)
    epp_tbl = doc.add_table(rows=n_rows + 1, cols=6)
    epp_tbl.style = "Table Grid"

    headers    = ["Elemento entregado", "Cantidad", "Fecha de Entrega", "Entregado por", "Recibí conforme", "Observación"]
    col_widths = [Cm(5.5), Cm(2.0), Cm(3.0), Cm(3.0), Cm(3.0), Cm(3.0)]
    for i, h in enumerate(headers):
        _cell_text(epp_tbl.rows[0].cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg_epp(epp_tbl.rows[0].cells[i], "D9EAD3")

    fecha_str = f"{fecha_entrega.day:02d}-{fecha_entrega.month:02d}-{fecha_entrega.year}" if fecha_entrega else ""
    for idx in range(n_rows):
        row = epp_tbl.rows[idx + 1]
        row.height = Cm(0.7)
        if idx < len(items):
            item = items[idx]
            _cell_text(row.cells[0], item.get("elemento", item.get("item", "")), size=8)
            _cell_text(row.cells[1], str(item.get("cantidad", 1)), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[2], fecha_str, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row.cells[3], entregado_por, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in epp_tbl.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    doc.add_paragraph()

    # ── Declaraciones ─────────────────────────────────────────────────────────
    for txt in [
        "El trabajador se compromete a mantener los Elementos de Protección Personal en buen estado, "
        "declara haberlos recibido en forma gratuita y usarlos correctamente cada vez que una actividad lo requiera.",
        "El trabajador declara que al momento de recibir sus elementos de protección personal ha recibido la "
        "capacitación necesaria para el correcto uso y cuidado de estos.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(txt)
        r.font.size = Pt(8); r.italic = True

    # ── Firmas ────────────────────────────────────────────────────────────────
    firma = doc.add_table(rows=2, cols=2)
    firma.style = "Table Grid"
    _cell_text(firma.rows[0].cells[0], "Entregado por", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[0].cells[1], "Firma Trabajador / Recibí conforme", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[0], entregado_por, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in firma.rows:
        row.cells[0].width = Cm(9.75)
        row.cells[1].width = Cm(9.75)
        row.height = Cm(2.0)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generar_pacto_horas_extra_docx(empresa, empleado, contrato, pacto, cargo_nombre: str = "") -> bytes:
    from decimal import Decimal
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Logo
    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))

    _parrafo(doc, ["PACTO DE HORAS EXTRAORDINARIAS"], bold_default=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _parrafo(doc, ["(Artículo 32 del Código del Trabajo)"], bold_default=False,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    nombre_completo = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()
    recargo_pct = int(Decimal(str(pacto.porcentaje_recargo)) * 100)

    _parrafo(doc, [
        f"En {empresa.ciudad or 'Santiago'}, a ", (_fecha_larga(pacto.fecha_inicio), True),
        ", entre ", (empresa.razon_social or "", True),
        ", RUT N° ", (empresa.rut or "", True),
        ", representado por ", (empresa.representante_legal or "", True),
        ', en adelante "EL EMPLEADOR", y don/doña ',
        (nombre_completo, True), ", RUT ", (empleado.rut or "", True),
        f', que se desempeña como {cargo_nombre or "trabajador/a"}, en adelante "EL TRABAJADOR", se conviene el siguiente Pacto de Horas Extraordinarias:',
    ])

    _parrafo(doc, [
        ("PRIMERO: ", True),
        "Las partes acuerdan que EL TRABAJADOR podrá laborar horas extraordinarias, con un tope de ",
        (f"{pacto.tope_horas_diarias}", True), " horas diarias adicionales a la jornada ordinaria pactada.",
    ])

    _parrafo(doc, [
        ("SEGUNDO: ", True),
        "Las horas extraordinarias pactadas se pagarán con un recargo del ",
        (f"{recargo_pct}%", True), " sobre el valor de la hora ordinaria, conforme a lo establecido en el artículo 32 del Código del Trabajo.",
    ])

    _parrafo(doc, [
        ("TERCERO: ", True),
        "El presente pacto regirá desde el ",
        (_fecha_larga(pacto.fecha_inicio), True), " hasta el ",
        (_fecha_larga(pacto.fecha_termino), True),
        ", pudiendo ser renovado por acuerdo de las partes.",
    ])

    _parrafo(doc, [
        ("CUARTO: ", True),
        "Las horas extraordinarias sólo podrán pactarse para atender necesidades o situaciones temporales de la empresa. "
        "Transcurrido el período de vigencia, las horas que se laboren en exceso de la jornada pactada no se considerarán extraordinarias, "
        "salvo que se suscriba un nuevo pacto.",
    ])

    _parrafo(doc, ["En señal de conformidad, las partes firman el presente pacto en dos ejemplares de igual valor y fecha."],
             space_after=40)

    # Firmas
    firma = doc.add_table(rows=2, cols=2)
    firma.style = "Table Grid"
    _cell_text(firma.rows[0].cells[0], "\n\n\n____________________________\nFirma Empleador",
               size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[0].cells[1], "\n\n\n____________________________\nFirma Trabajador",
               size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[0], empresa.razon_social or "", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[1], nombre_completo, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in firma.rows:
        row.cells[0].width = Cm(9.75)
        row.cells[1].width = Cm(9.75)
        row.height = Cm(2.0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Causales legales de despido ──────────────────────────────────────────────
CAUSALES_DESPIDO = {
    # Art. 159 – Término sin responsabilidad del empleador
    "159_1": ("Art. 159 N°1", "Mutuo acuerdo de las partes", False, False),
    "159_2": ("Art. 159 N°2", "Renuncia del trabajador", False, False),
    "159_4": ("Art. 159 N°4", "Vencimiento del plazo convenido en el contrato", False, False),
    "159_5": ("Art. 159 N°5", "Conclusión del trabajo o servicio que dio origen al contrato", False, False),
    "159_6": ("Art. 159 N°6", "Caso fortuito o fuerza mayor", False, False),
    # Art. 160 – Despido disciplinario (sin indemnización)
    "160_1": ("Art. 160 N°1", "Falta de probidad del trabajador en el desempeño de sus funciones", False, False),
    "160_1b": ("Art. 160 N°1 letra b)", "Conductas de acoso sexual", False, False),
    "160_1f": ("Art. 160 N°1 letra f)", "Conductas de acoso laboral", False, False),
    "160_3": ("Art. 160 N°3", "No concurrencia del trabajador a sus labores sin causa justificada durante dos días seguidos, dos lunes en el mes o un total de tres días en el mes", False, False),
    "160_4": ("Art. 160 N°4", "Abandono del trabajo por parte del trabajador", False, False),
    "160_5": ("Art. 160 N°5", "Actos, omisiones o imprudencias temerarias que afecten a la seguridad o al funcionamiento del establecimiento, a la seguridad o a la actividad de los trabajadores, o a la salud de éstos", False, False),
    "160_7": ("Art. 160 N°7", "Incumplimiento grave de las obligaciones que impone el contrato de trabajo", False, False),
    # Art. 161 – Necesidades de la empresa (con indemnización)
    "161_1": ("Art. 161 inciso 1°", "Necesidades de la empresa, establecimiento o servicio, tales como las derivadas de la racionalización o modernización de los mismos, bajas en la productividad, cambios en las condiciones del mercado o de la economía, que hagan necesaria la separación de uno o más trabajadores", True, True),
    "161_2": ("Art. 161 inciso 2°", "Desahucio del empleador (trabajadores con poder de representación o de dirección, o trabajadores de casa particular)", True, True),
}

MOTIVOS_AMONESTACION = [
    "Atrasos reiterados e injustificados al lugar de trabajo",
    "Ausencia injustificada al trabajo",
    "Incumplimiento del Reglamento Interno de Orden, Higiene y Seguridad",
    "Trato irrespetuoso hacia compañeros de trabajo o jefaturas",
    "Negligencia o descuido en el desempeño de sus funciones",
    "No uso de Elementos de Protección Personal (EPP) obligatorios",
    "Uso indebido de herramientas, equipos o bienes de la empresa",
    "Incumplimiento de instrucciones impartidas por el empleador",
    "Conducta inapropiada en el lugar de trabajo",
    "Daño a bienes de la empresa por dolo o negligencia",
    "Otro motivo",
]


def generar_amonestacion_docx(empresa, empleado, motivo: str, descripcion: str, fecha: date, cargo_nombre: str = "") -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))

    _parrafo(doc, ["CARTA DE AMONESTACIÓN"], bold_default=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _parrafo(doc, [f"{empresa.ciudad or 'Santiago'}, {_fecha_larga(fecha)}"],
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=16)

    nombre_completo = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()

    _parrafo(doc, [
        "Señor(a)\n", (nombre_completo, True), "\n",
        f"Cargo: {cargo_nombre or 'Trabajador/a'}\n",
        f"RUT: {empleado.rut or ''}\n",
        "Presente",
    ], space_after=16, align=WD_ALIGN_PARAGRAPH.LEFT)

    _parrafo(doc, ["Estimado/a señor(a) ", (nombre_completo, True), ":"], space_after=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    _parrafo(doc, [
        "Por medio de la presente, ", (empresa.razon_social or "", True),
        f", RUT {empresa.rut or ''}, se dirige a usted con el objeto de comunicarle formalmente ",
        ("una amonestación escrita", True),
        f", en razón de los siguientes hechos:",
    ])

    _parrafo(doc, [motivo.upper()], bold_default=True, space_after=6)
    if descripcion:
        _parrafo(doc, [descripcion], space_after=14)

    _parrafo(doc, [
        "Lo anterior constituye un incumplimiento a las obligaciones que le impone el contrato de trabajo y el Reglamento Interno de Orden, Higiene y Seguridad, "
        "y podría ser constitutivo de causal de término de contrato conforme al artículo 160 N°7 del Código del Trabajo en caso de reincidencia."
    ])

    _parrafo(doc, [
        "Le solicitamos enmiende su conducta para evitar medidas más graves. La presente carta será archivada en su carpeta de personal.",
    ], space_after=40)

    # Firmas
    firma = doc.add_table(rows=2, cols=2)
    firma.style = "Table Grid"
    _cell_text(firma.rows[0].cells[0], "\n\n\n____________________________\nFirma Empleador",
               size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[0].cells[1], "\n\n\n____________________________\nFirma Trabajador / Notificado",
               size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[0], empresa.razon_social or "", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[1], nombre_completo, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in firma.rows:
        row.cells[0].width = Cm(9.75)
        row.cells[1].width = Cm(9.75)
        row.height = Cm(2.0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generar_carta_despido_docx(
    empresa, empleado, contrato,
    causal_codigo: str,
    fecha_termino: date,
    cargo_nombre: str = "",
    dias_trabajados_mes: int = 0,
    monto_dias_trabajados: int = 0,
    vacaciones_proporcionales: int = 0,
    indemnizacion_anos: int = 0,
    aviso_previo: int = 0,
    anos_servicio: int = 0,
    gratificacion: int = 0,
    rem_pendiente: int = 0,
    descripcion_adicional: str = "",
) -> bytes:
    from decimal import Decimal
    causal_info = CAUSALES_DESPIDO.get(causal_codigo, ("", causal_codigo, False, False))
    art_ref, causal_texto, tiene_indem, tiene_aviso = causal_info

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    logo = _logo_bytes(getattr(empresa, "logo_url", None))
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(io.BytesIO(logo), height=Cm(1.6))

    _parrafo(doc, ["CARTA DE AVISO DE TÉRMINO DE CONTRATO DE TRABAJO"], bold_default=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _parrafo(doc, [f"{empresa.ciudad or 'Santiago'}, {_fecha_larga(fecha_termino)}"],
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=16)

    nombre_completo = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()

    _parrafo(doc, [
        "Señor(a)\n", (nombre_completo, True), "\n",
        f"Cargo: {cargo_nombre or 'Trabajador/a'}\n",
        f"RUT: {empleado.rut or ''}\n",
        "Presente",
    ], space_after=16, align=WD_ALIGN_PARAGRAPH.LEFT)

    _parrafo(doc, ["Estimado/a señor(a) ", (nombre_completo, True), ":"], space_after=10, align=WD_ALIGN_PARAGRAPH.LEFT)

    _parrafo(doc, [
        f"Por medio de la presente, {empresa.razon_social or ''}, RUT {empresa.rut or ''}, ",
        "comunica a usted el término de su contrato de trabajo a contar del día ",
        (_fecha_larga(fecha_termino), True),
        f", invocando la causal contemplada en el ", (art_ref, True),
        " del Código del Trabajo, esto es: ",
    ])

    _parrafo(doc, [f'"{causal_texto}."'], bold_default=True, space_after=10)

    if descripcion_adicional:
        _parrafo(doc, [descripcion_adicional], space_after=10)

    if contrato.fecha_inicio:
        _parrafo(doc, [
            f"El trabajador ingresó a prestar servicios el día ",
            (_fecha_larga(contrato.fecha_inicio), True),
            f", acumulando {anos_servicio} año(s) de servicios.",
        ], space_after=10)

    # Tabla de montos
    total = monto_dias_trabajados + vacaciones_proporcionales + indemnizacion_anos + aviso_previo + gratificacion + rem_pendiente
    items = [
        ("Remuneración días trabajados del mes", monto_dias_trabajados),
    ]
    if rem_pendiente > 0:
        items.append((f"Colación y movilización proporcional — {dias_trabajados_mes} días", rem_pendiente))
    items.append(("Vacaciones proporcionales (Art. 67 CT)", vacaciones_proporcionales))
    if gratificacion > 0:
        items.append((f"Gratificación proporcional (Art. 50 CT) — {dias_trabajados_mes} días", gratificacion))
    if tiene_indem and indemnizacion_anos > 0:
        items.append((f"Indemnización por años de servicio (Art. 163 CT) — {anos_servicio} año(s)", indemnizacion_anos))
    if tiene_aviso and aviso_previo > 0:
        items.append(("Indemnización sustitutiva de aviso previo (Art. 161 CT)", aviso_previo))

    _parrafo(doc, ["En virtud de lo anterior, se pone a su disposición el siguiente finiquito:"], space_after=8)

    tbl = doc.add_table(rows=len(items) + 2, cols=2)
    tbl.style = "Table Grid"
    # Header
    _cell_text(tbl.rows[0].cells[0], "CONCEPTO", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(tbl.rows[0].cells[1], "MONTO ($)", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, (concepto, monto) in enumerate(items, 1):
        _cell_text(tbl.rows[i].cells[0], concepto, size=10)
        _cell_text(tbl.rows[i].cells[1], f"${monto:,}".replace(",", "."), size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
    # Total
    total_row = tbl.rows[len(items) + 1]
    _cell_text(total_row.cells[0], "TOTAL FINIQUITO", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _cell_text(total_row.cells[1], f"${total:,}".replace(",", "."), bold=True, size=11, align=WD_ALIGN_PARAGRAPH.RIGHT)
    tbl.columns[0].width = Cm(12.5)
    tbl.columns[1].width = Cm(5.0)

    doc.add_paragraph()
    _parrafo(doc, [
        "El trabajador tiene derecho a reclamar ante la Inspección del Trabajo o los Tribunales de Justicia "
        "dentro del plazo de 60 días hábiles desde la separación (Art. 168 CT). El finiquito tendrá poder "
        "liberatorio cuando sea ratificado ante ministro de fe (Notario, Inspector del Trabajo o presidente "
        "de la organización sindical, Art. 177 CT).",
    ], space_after=40)

    # Firmas
    firma = doc.add_table(rows=2, cols=2)
    firma.style = "Table Grid"
    _cell_text(firma.rows[0].cells[0], "\n\n\n____________________________\nFirma Empleador",
               size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[0].cells[1], "\n\n\n____________________________\nFirma Trabajador / Recibí conforme",
               size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[0], empresa.razon_social or "", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(firma.rows[1].cells[1], nombre_completo, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in firma.rows:
        row.cells[0].width = Cm(9.75)
        row.cells[1].width = Cm(9.75)
        row.height = Cm(2.0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
