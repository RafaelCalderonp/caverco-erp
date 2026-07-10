"""Generadores de documentos Word para módulo de documentación de obra."""
import base64
import io
import os
from datetime import date, datetime
from decimal import Decimal

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LOGO_PATH = os.path.join(os.path.dirname(__file__), "..", "assets", "archimet_logo.png")

MOTIVOS = {
    "CHARLA_INDUCCION": "Charla de Inducción",
    "CAPACITACION":     "Capacitación",
    "ENTRENAMIENTO":    "Entrenamientos",
    "OTRAS":            "Otras",
}

MESES = ["enero","febrero","marzo","abril","mayo","junio",
         "julio","agosto","septiembre","octubre","noviembre","diciembre"]


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_borders(cell, color="000000", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _cell_para(cell, text="", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
               color=None, italic=False, space_before=1, space_after=1):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(str(text))
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p


def _checkbox_cell(cell, checked: bool, label: str, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label)
    run.font.size = Pt(size)
    # Segunda línea con el checkbox
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(0)
    r2 = p2.add_run("☑" if checked else "☐")
    r2.font.size = Pt(10)


def _logo_bytes_from_url(logo_url: str | None) -> bytes | None:
    if not logo_url or not logo_url.startswith("data:"):
        return None
    try:
        return base64.b64decode(logo_url.split(",", 1)[1])
    except Exception:
        return None


def _add_logo(cell, logo_url: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    logo_data = _logo_bytes_from_url(logo_url)
    if logo_data:
        run.add_picture(io.BytesIO(logo_data), width=Cm(2.8))
    elif os.path.exists(LOGO_PATH):
        run.add_picture(LOGO_PATH, width=Cm(2.8))
    else:
        run.text = "ARCHIMET"
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1F, 0x39, 0x6F)


def _fmt_date(d) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        return d
    return f"{d.day:02d}-{d.month:02d}-{d.year}"


def _fmt_date_larga(d) -> str:
    if not d:
        return ""
    if isinstance(d, str):
        return d
    return f"{d.day:02d} {MESES[d.month-1]} {d.year}"


def _new_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21.59)
    section.page_height   = Cm(27.94)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    return doc


# ─── 1. REGISTRO DE ASISTENCIA A CAPACITACIÓN ────────────────────────────────

def generar_capacitacion_docx(capacitacion, procedimiento, asistentes, empresa_nombre: str, empresa=None) -> bytes:
    doc = _new_doc()

    obj_general   = capacitacion.objetivo_general   or (procedimiento.objetivo_general   if procedimiento else "") or ""
    obj_especifico = capacitacion.objetivos_especificos or (procedimiento.objetivos_especificos if procedimiento else "") or ""

    # ── Encabezado ──────────────────────────────────────────────────────────
    hdr = doc.add_table(rows=2, cols=3)
    hdr.style = "Table Grid"
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Logo (merge rows)
    logo_cell = hdr.cell(0, 0)
    logo_cell.merge(hdr.cell(1, 0))
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _add_logo(logo_cell, getattr(empresa, "logo_url", None))

    # Título centro
    _cell_para(hdr.cell(0, 1), "Documento de Prevención de Riesgos.", bold=False, size=9,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr.cell(1, 1), "Registro de Asistencia a Capacitación.", bold=True, size=10,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    # Fechas derecha (versión + fecha modificación)
    _cell_para(hdr.cell(0, 2), "4-11-2025", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr.cell(1, 2), "17-11-2025", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    for row in hdr.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(10.5)
        row.cells[2].width = Cm(3.3)

    doc.add_paragraph()

    # ── Fila versión / fecha / hora ──────────────────────────────────────────
    info1 = doc.add_table(rows=3, cols=4)
    info1.style = "Table Grid"

    data_info1 = [
        ("Versión:", capacitacion.version or "01", "Versión:", capacitacion.version or "01"),
        ("Fecha:", _fmt_date(capacitacion.fecha), "Hora de Inicio:", capacitacion.hora_inicio or ""),
        ("Duración:", f"{capacitacion.duracion_horas or 1} h", "Hora de Término:", capacitacion.hora_termino or ""),
    ]

    for i, (l1, v1, l2, v2) in enumerate(data_info1):
        _cell_para(info1.rows[i].cells[0], l1, bold=True, size=8)
        _cell_para(info1.rows[i].cells[1], v1, size=8)
        _cell_para(info1.rows[i].cells[2], l2, bold=True, size=8)
        _cell_para(info1.rows[i].cells[3], v2, size=8)

    # Ajustar: ocultar la celda versión duplicada — simplificamos usando la misma tabla
    # Fila 0: Versión: 01  | (vacío) | (vacío) | (vacío)
    # Rewrite row 0 properly
    _cell_para(info1.rows[0].cells[0], "Versión:", bold=True, size=8)
    _cell_para(info1.rows[0].cells[1], capacitacion.version or "01", size=8)
    _cell_para(info1.rows[0].cells[2], "Hora de Inicio:", bold=True, size=8)
    _cell_para(info1.rows[0].cells[3], capacitacion.hora_inicio or "8:00", size=8)

    _cell_para(info1.rows[1].cells[0], "Fecha:", bold=True, size=8)
    _cell_para(info1.rows[1].cells[1], _fmt_date(capacitacion.fecha), size=8)
    _cell_para(info1.rows[1].cells[2], "Hora de Término:", bold=True, size=8)
    _cell_para(info1.rows[1].cells[3], capacitacion.hora_termino or "9:00", size=8)

    _cell_para(info1.rows[2].cells[0], "Duración:", bold=True, size=8)
    _cell_para(info1.rows[2].cells[1], f"{capacitacion.duracion_horas or 1} h", size=8)
    # Fila 2 col 2-3: Motivo (checkboxes)
    motivo_cell = info1.rows[2].cells[2]
    motivo_cell.merge(info1.rows[2].cells[3])
    motivo_cell.text = ""
    mp = motivo_cell.paragraphs[0]
    mp.paragraph_format.space_before = Pt(1)
    mp.paragraph_format.space_after  = Pt(1)
    r_motivo = mp.add_run("Motivo:  ")
    r_motivo.bold = True
    r_motivo.font.size = Pt(8)

    for key, label in MOTIVOS.items():
        checked = (capacitacion.motivo == key)
        r_box = mp.add_run(("☑ " if checked else "☐ ") + label + "   ")
        r_box.font.size = Pt(8)

    for col_w in [2.5, 3.5, 3.5, 7.5]:
        pass  # widths set below
    col_widths1 = [2.5, 3.5, 3.5, 7.5]
    for row in info1.rows:
        for i, w in enumerate(col_widths1):
            row.cells[i].width = Cm(w)

    doc.add_paragraph()

    # ── Objetivos ────────────────────────────────────────────────────────────
    obj_table = doc.add_table(rows=2, cols=2)
    obj_table.style = "Table Grid"

    _cell_para(obj_table.rows[0].cells[0], "Objetivo General", bold=True, size=8)
    _cell_para(obj_table.rows[0].cells[1], obj_general, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    obj_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _cell_para(obj_table.rows[1].cells[0], "Objetivos Específicos", bold=True, size=8)

    # Objetivos específicos con viñetas (cada línea del texto)
    esp_cell = obj_table.rows[1].cells[1]
    esp_cell.text = ""
    lines = [l for l in (obj_especifico or "").split("\n") if l.strip()]
    for j, line in enumerate(lines):
        p = esp_cell.paragraphs[0] if j == 0 else esp_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line.strip())
        r.font.size = Pt(8)

    for row in obj_table.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(13.0)
    obj_table.rows[0].height = Cm(1.5)
    obj_table.rows[1].height = Cm(3.0)

    doc.add_paragraph()

    # ── Datos del Relator ────────────────────────────────────────────────────
    rel_hdr = doc.add_table(rows=1, cols=1)
    rel_hdr.style = "Table Grid"
    _cell_para(rel_hdr.rows[0].cells[0],
               "Datos de la Persona que Realizó la Capacitación.",
               bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_bg(rel_hdr.rows[0].cells[0], "D9EAD3")

    rel_table = doc.add_table(rows=2, cols=4)
    rel_table.style = "Table Grid"

    _cell_para(rel_table.rows[0].cells[0], "Nombre Completo", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(rel_table.rows[0].cells[1], "Área",            bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(rel_table.rows[0].cells[2], "RUT",             bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(rel_table.rows[0].cells[3], "Firma",           bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_bg(rel_table.rows[0].cells[0], "D9EAD3")
    _set_cell_bg(rel_table.rows[0].cells[1], "D9EAD3")
    _set_cell_bg(rel_table.rows[0].cells[2], "D9EAD3")
    _set_cell_bg(rel_table.rows[0].cells[3], "D9EAD3")

    _cell_para(rel_table.rows[1].cells[0], capacitacion.relator_nombre or "Salvador Calderón", size=8)
    _cell_para(rel_table.rows[1].cells[1], capacitacion.relator_area   or "Prevención de riesgos", size=8)
    _cell_para(rel_table.rows[1].cells[2], capacitacion.relator_rut    or "18.512.365-0", size=8)
    _cell_para(rel_table.rows[1].cells[3], "", size=8)
    rel_table.rows[1].height = Cm(1.0)

    rel_widths = [5.5, 4.0, 3.0, 4.5]
    for tbl in [rel_table]:
        for row in tbl.rows:
            for i, w in enumerate(rel_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()

    # ── Tabla de asistentes ──────────────────────────────────────────────────
    n_rows = max(len(asistentes), 15)
    att = doc.add_table(rows=n_rows + 1, cols=4)
    att.style = "Table Grid"

    for i, h in enumerate(["N°", "Nombre Completo", "Área", "RUT"]):
        _cell_para(att.rows[0].cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(att.rows[0].cells[i], "D9EAD3")

    for idx in range(n_rows):
        row = att.rows[idx + 1]
        _cell_para(row.cells[0], str(idx + 1), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        if idx < len(asistentes):
            a = asistentes[idx]
            _cell_para(row.cells[1], a.nombre or "", size=8)
            _cell_para(row.cells[2], a.area   or "", size=8)
            _cell_para(row.cells[3], a.rut    or "", size=8)
        else:
            for c in [1, 2, 3]:
                _cell_para(row.cells[c], "", size=8)
        row.height = Cm(0.65)

    att_widths = [1.0, 7.5, 5.0, 3.5]
    for row in att.rows:
        for i, w in enumerate(att_widths):
            row.cells[i].width = Cm(w)

    # Footer
    doc.add_paragraph()
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(empresa_nombre)
    rf.font.size = Pt(7)
    rf.font.color.rgb = RGBColor(130, 130, 130)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 2. REGISTRO CAPACITACIÓN ARCHIMET (formato Instalaciones Arquitectónicas) ──

ARCHIMET_CATEGORIAS = [
    "CHARLA ESPECIFICA",
    "CHARLA OPERACIONAL",
    "CHARLA INTEGRAGRAL SEMANAL",
    "REINDUCCION",
    "CURSO DE CAPACITACION/ TALLER",
    "CONTACTO PERSONAL",
]


def generar_capacitacion_archimet_docx(capacitacion, procedimiento, asistentes, empresa=None) -> bytes:
    doc = _new_doc()

    n_asistentes = len(asistentes)
    duracion     = float(capacitacion.duracion_horas or 1)
    total_hh     = round(n_asistentes * duracion, 1)

    tema_titulo = (procedimiento.objetivo_general or "") if procedimiento else (capacitacion.objetivo_general or "")
    tema_cuerpo = (procedimiento.objetivos_especificos or "") if procedimiento else (capacitacion.objetivos_especificos or "")

    # ── Encabezado ──────────────────────────────────────────────────────────
    hdr = doc.add_table(rows=2, cols=3)
    hdr.style = "Table Grid"
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER

    logo_cell = hdr.cell(0, 0)
    logo_cell.merge(hdr.cell(1, 0))
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _add_logo(logo_cell, getattr(empresa, "logo_url", None))

    _cell_para(hdr.cell(0, 1), "SISTEMA DE GESTION INTEGRADOS", bold=False, size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr.cell(1, 1), "REGISTRO DE CAPACITACION Y ENTRENAMIENTO", bold=True, size=10,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    _cell_para(hdr.cell(0, 2), "REV: 02", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(hdr.cell(1, 2), f"FECHA: {capacitacion.fecha.year if hasattr(capacitacion.fecha, 'year') else '2026'}",
               bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    col_widths_hdr = [3.0, 11.0, 3.0]
    for row in hdr.rows:
        for i, w in enumerate(col_widths_hdr):
            row.cells[i].width = Cm(w)

    p_sel = doc.add_paragraph()
    p_sel.paragraph_format.space_before = Pt(4)
    p_sel.paragraph_format.space_after  = Pt(4)
    r_sel = p_sel.add_run('SELECCIONE CON UNA "X" EL ALCANCE DEL ENTRENAMIENTO O CAPACITACION')
    r_sel.bold = True
    r_sel.font.size = Pt(8)

    # ── Tabla categorías + resumen ──────────────────────────────────────────
    # 7 cols: [CATEGORIA] [SSO] [MA] [CAL.] [sep] [RIGHT_LABEL] [RIGHT_VALUE]
    # 7 rows: header + 6 categorías
    cat_tbl = doc.add_table(rows=7, cols=7)
    cat_tbl.style = "Table Grid"

    col_w = [5.2, 1.1, 1.1, 1.1, 0.1, 4.5, 3.9]

    # Headers row 0
    _cell_para(cat_tbl.rows[0].cells[0], "CATEGORIA", bold=True, size=8)
    _cell_para(cat_tbl.rows[0].cells[1], "SSO", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(cat_tbl.rows[0].cells[2], "MA",  bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(cat_tbl.rows[0].cells[3], "CAL.", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(cat_tbl.rows[0].cells[4], "", size=8)
    _cell_para(cat_tbl.rows[0].cells[5], "TOTAL PERSONAL ENTRENADO", bold=True, size=8)
    _cell_para(cat_tbl.rows[0].cells[6], str(n_asistentes), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    right_labels = [
        ("DURACION (horas)", str(duracion)),
        ("", ""),
        ("TOTAL HH", str(total_hh)),
        (f"FECHA:", _fmt_date(capacitacion.fecha)),
        (f"HORA:",  capacitacion.hora_inicio or ""),
        ("", ""),
    ]

    for i, cat in enumerate(ARCHIMET_CATEGORIAS):
        row = cat_tbl.rows[i + 1]
        _cell_para(row.cells[0], cat, size=8)
        # SSO column: X only for CHARLA ESPECIFICA
        sso_val = "X" if i == 0 else ""
        _cell_para(row.cells[1], sso_val, bold=(i == 0), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[2], "", size=8)
        _cell_para(row.cells[3], "", size=8)
        _cell_para(row.cells[4], "", size=8)
        lbl, val = right_labels[i]
        _cell_para(row.cells[5], lbl, bold=bool(lbl), size=8)
        _cell_para(row.cells[6], val, size=8)

    for row in cat_tbl.rows:
        for i, w in enumerate(col_w):
            row.cells[i].width = Cm(w)
        row.height = Cm(0.55)

    # ── Tema Tratado ────────────────────────────────────────────────────────
    tema_tbl = doc.add_table(rows=1, cols=2)
    tema_tbl.style = "Table Grid"

    tema_label_cell = tema_tbl.rows[0].cells[0]
    tema_label_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _cell_para(tema_label_cell, "TEMA TRATADO :", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    tema_label_cell.width = Cm(3.0)

    tema_cell = tema_tbl.rows[0].cells[1]
    tema_cell.text = ""

    # Primera línea (título del tema) en negrita
    p0 = tema_cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after  = Pt(2)
    r0 = p0.add_run(tema_titulo)
    r0.bold = True
    r0.font.size = Pt(8)

    # Líneas del cuerpo (puntos o texto)
    for line in tema_cuerpo.split("\n"):
        line = line.strip()
        if not line:
            continue
        pb = tema_cell.add_paragraph()
        pb.paragraph_format.space_before = Pt(1)
        pb.paragraph_format.space_after  = Pt(1)
        rb = pb.add_run(line)
        rb.font.size = Pt(8)

    tema_cell.width = Cm(14.0)
    tema_tbl.rows[0].height = Cm(3.5)

    # ── Datos de la capacitación ────────────────────────────────────────────
    datos_tbl = doc.add_table(rows=5, cols=4)
    datos_tbl.style = "Table Grid"

    # Row 0: OBRA
    obra_cell = datos_tbl.rows[0].cells[0]
    obra_cell.merge(datos_tbl.rows[0].cells[3])
    _cell_para(obra_cell, f"OBRA:  {capacitacion.obra or ''}", bold=False, size=8)

    # Row 1: NOMBRE DEL RELATOR
    rel_cell = datos_tbl.rows[1].cells[0]
    rel_cell.merge(datos_tbl.rows[1].cells[3])
    _cell_para(rel_cell, f"NOMBRE DEL RELATOR :  {capacitacion.relator_nombre or ''}", size=8)

    # Row 2: CARGO + FIRMA (split)
    datos_tbl.rows[2].cells[0].merge(datos_tbl.rows[2].cells[1])
    _cell_para(datos_tbl.rows[2].cells[0],
               f"CARGO:  {capacitacion.relator_area or ''}", size=8)
    datos_tbl.rows[2].cells[2].merge(datos_tbl.rows[2].cells[3])
    _cell_para(datos_tbl.rows[2].cells[2], "FIRMA :", bold=True, size=8)

    # Row 3: LUGAR
    lug_cell = datos_tbl.rows[3].cells[0]
    lug_cell.merge(datos_tbl.rows[3].cells[3])
    _cell_para(lug_cell,
               f"LUGAR / ESTABLECIMIENTO :  {capacitacion.lugar_establecimiento or ''}", size=8)

    # Row 4: MATERIAL DE APOYO
    mat_cell = datos_tbl.rows[4].cells[0]
    mat_cell.merge(datos_tbl.rows[4].cells[3])
    _cell_para(mat_cell,
               f"MATERIAL DE APOYO :  {capacitacion.material_apoyo or ''}", size=8)

    col_w_datos = [4.0, 4.0, 4.0, 5.0]
    for row in datos_tbl.rows:
        for i, w in enumerate(col_w_datos):
            row.cells[i].width = Cm(w)
        row.height = Cm(0.7)
    datos_tbl.rows[2].height = Cm(1.2)

    # ── Tabla de participantes ───────────────────────────────────────────────
    n_rows = max(len(asistentes), 7)
    part_tbl = doc.add_table(rows=n_rows + 1, cols=5)
    part_tbl.style = "Table Grid"

    headers = ["Nº", "NOMBRE", "CARGO", "RUT", "FIRMA"]
    col_w_part = [1.0, 5.0, 3.5, 3.0, 4.5]

    for i, h in enumerate(headers):
        _cell_para(part_tbl.rows[0].cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(part_tbl.rows[0].cells[i], "D9D9D9")

    for idx in range(n_rows):
        row = part_tbl.rows[idx + 1]
        _cell_para(row.cells[0], f"{idx + 1} -", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        if idx < len(asistentes):
            a = asistentes[idx]
            _cell_para(row.cells[1], a.nombre or "", size=8)
            _cell_para(row.cells[2], a.area   or "", size=8)
            _cell_para(row.cells[3], a.rut    or "", size=8)
        else:
            for c in [1, 2, 3]:
                _cell_para(row.cells[c], "", size=8)
        _cell_para(row.cells[4], "", size=8)
        row.height = Cm(0.8)

    for row in part_tbl.rows:
        for i, w in enumerate(col_w_part):
            row.cells[i].width = Cm(w)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 2. ENTREGA DE REGLAMENTO INTERNO ────────────────────────────────────────

def generar_reglamento_docx(nombre: str, rut: str, seccion: str,
                             fecha: date, empresa_nombre: str) -> bytes:
    doc = _new_doc()

    # Título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(16)
    r = p_titulo.add_run(
        "FORMULARIO RECEPCIÓN DEL REGLAMENTO INTERNO DE ORDEN, HIGIENE Y SEGURIDAD, "
        "ART. 67º DE LA LEY Nº 16.744, TITULO III DEL CÓDIGO DEL TRABAJO, D.F.L. N° 1"
    )
    r.bold = True
    r.font.size = Pt(11)

    # Párrafos declaración
    textos = [
        f"Declaro haber recibido en forma gratuita una copia del reglamento interno de orden, higiene y seguridad "
        f"de la empresa {empresa_nombre} de acuerdo a lo establecido en el Art. 56 del D.S N° 40 y Art. 67 de la Ley 16.744.",
        "Asumo mi responsabilidad de dar lectura a su contenido y dar cumplimiento a las obligaciones, prohibiciones, "
        "normas de orden, higiene y seguridad que en él están escritas, como así también a las disposiciones "
        "establecidas por el Organismo Administrador del Seguro de Accidentes del Trabajo y Enfermedades "
        "Profesionales a que se encuentre adherida la empresa.",
    ]
    for txt in textos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(txt)
        r.font.size = Pt(10)

    # Separador
    p_sep = doc.add_paragraph("- " * 60)
    p_sep.paragraph_format.space_after = Pt(16)
    p_sep.runs[0].font.size = Pt(8)

    # Tabla de datos
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"

    rows_data = [
        ("Nombre completo", nombre or ""),
        ("R.U.T.",          rut    or ""),
        ("Sección",         seccion or ""),
        ("Firma del trabajador", ""),
        ("Fecha de entrega", _fmt_date_larga(fecha) if fecha else ""),
    ]

    for i, (label, valor) in enumerate(rows_data):
        _cell_para(tbl.rows[i].cells[0], label, bold=True, size=10)
        _cell_para(tbl.rows[i].cells[1], valor, size=10)
        tbl.rows[i].cells[0].width = Cm(5.5)
        tbl.rows[i].cells[1].width = Cm(11.5)
        if label == "Firma del trabajador":
            tbl.rows[i].height = Cm(2.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 3. ENTREGA DE EPP ───────────────────────────────────────────────────────

EPP_DEFAULT = [
    "Casco", "Zapatos de Seguridad", "Guantes", "Antiparras/Lentes",
    "Tapones Auditivos", "Barbiquejo", "Bloqueador Solar",
    "Chaleco Reflectante", "Arnés de Seguridad", "Cabo de vida simple",
]


def generar_epp_docx(nombre: str, rut: str, cargo: str, obra: str,
                     fecha: date, items: list, entregado_por: str,
                     empresa_nombre: str) -> bytes:
    doc = _new_doc()

    # Encabezado logo + título
    hdr = doc.add_table(rows=1, cols=2)
    hdr.style = "Table Grid"
    _add_logo(hdr.rows[0].cells[0])
    _cell_para(hdr.rows[0].cells[1],
               "REGISTRO DE ENTREGA DE ELEMENTOS DE PROTECCIÓN PERSONAL",
               bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    hdr.rows[0].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    hdr.rows[0].cells[0].width = Cm(3.5)
    hdr.rows[0].cells[1].width = Cm(13.5)

    # Texto legal
    doc.add_paragraph()
    p_ley = doc.add_paragraph()
    p_ley.paragraph_format.space_after = Pt(8)
    r_ley = p_ley.add_run(
        'De acuerdo a lo estipulado en la Ley 16.744 Art. 68 inciso tres "Las empresas deberán proporcionar a sus '
        'trabajadores los equipos e implementos de protección necesarios no pudiendo en caso alguno cobrarles su valor".'
    )
    r_ley.font.size = Pt(8)
    r_ley.italic = True

    # Datos del trabajador
    datos_tbl = doc.add_table(rows=4, cols=2)
    datos_tbl.style = "Table Grid"
    datos_rows = [
        ("Nombre del trabajador", nombre or ""),
        ("Cédula de Identidad",   rut    or ""),
        ("Cargo",                 cargo  or ""),
        ("Área / Obra / Faena",   obra   or empresa_nombre),
    ]
    for i, (lbl, val) in enumerate(datos_rows):
        _cell_para(datos_tbl.rows[i].cells[0], lbl, bold=True, size=8)
        _cell_para(datos_tbl.rows[i].cells[1], val, size=8)
        datos_tbl.rows[i].cells[0].width = Cm(5.0)
        datos_tbl.rows[i].cells[1].width = Cm(12.0)

    doc.add_paragraph()

    # Subtítulo EPP
    p_sub = doc.add_paragraph("REGISTRO PERSONAL DE ENTREGA DE EPP")
    p_sub.runs[0].bold = True
    p_sub.runs[0].font.size = Pt(9)
    p_sub.paragraph_format.space_after = Pt(4)

    # Tabla EPP
    epp_tbl = doc.add_table(rows=len(items) + 1, cols=4)
    epp_tbl.style = "Table Grid"

    for i, h in enumerate(["Elemento entregado", "Cantidad", "Fecha de Entrega", "Recibí conforme / Observación"]):
        _cell_para(epp_tbl.rows[0].cells[i], h, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(epp_tbl.rows[0].cells[i], "D9EAD3")

    for idx, item in enumerate(items):
        row = epp_tbl.rows[idx + 1]
        _cell_para(row.cells[0], item.get("elemento", ""), size=8)
        _cell_para(row.cells[1], str(item.get("cantidad", 1)), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[2], _fmt_date(item.get("fecha") or fecha), size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[3], "", size=8)
        row.height = Cm(0.7)

    epp_widths = [6.0, 2.0, 3.5, 5.5]
    for row in epp_tbl.rows:
        for i, w in enumerate(epp_widths):
            row.cells[i].width = Cm(w)

    doc.add_paragraph()

    # Declaraciones
    declaraciones = [
        "El trabajador se compromete a mantener los Elementos de Protección Personal en buen estado, declara haberlos "
        "recibido en forma gratuita y usarlos correctamente cada vez que una actividad lo requiera.",
        "El trabajador declara que al momento de recibir sus elementos de protección personal ha recibido la "
        "capacitación necesaria para el correcto uso y cuidado de estos.",
    ]
    for txt in declaraciones:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(txt)
        r.font.size = Pt(8)
        r.italic = True

    # Firma
    firma_tbl = doc.add_table(rows=2, cols=2)
    firma_tbl.style = "Table Grid"
    _cell_para(firma_tbl.rows[0].cells[0], "Entregado por", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(firma_tbl.rows[0].cells[1], "Firma Trabajador", bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(firma_tbl.rows[1].cells[0], entregado_por or "Salvador Calderón", size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(firma_tbl.rows[1].cells[1], "", size=8)
    for row in firma_tbl.rows:
        row.cells[0].width = Cm(8.5)
        row.cells[1].width = Cm(8.5)
        row.height = Cm(2.0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 4. CERTIFICADO DE ANTIGÜEDAD ────────────────────────────────────────────

def generar_certificado_antiguedad_docx(
    nombre: str, rut_empleado: str, cargo: str,
    fecha_ingreso: date, tipo_contrato: str,
    empresa_nombre: str, empresa_rut: str,
    ciudad: str = "Santiago",
    fecha_emision: date = None,
    empresa=None,
) -> bytes:
    doc = _new_doc()
    if fecha_emision is None:
        fecha_emision = date.today()

    # Logo de la empresa
    logo_url = getattr(empresa, "logo_url", None) if empresa else None
    logo_data = _logo_bytes_from_url(logo_url)
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_after = Pt(12)
    run_logo = p_logo.add_run()
    if logo_data:
        run_logo.add_picture(io.BytesIO(logo_data), width=Cm(4))
    elif os.path.exists(LOGO_PATH):
        run_logo.add_picture(LOGO_PATH, width=Cm(4))

    doc.add_paragraph()

    # Título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.paragraph_format.space_after = Pt(30)
    r = p_titulo.add_run("CERTIFICADO DE ANTIGÜEDAD")
    r.bold = True
    r.font.size = Pt(14)

    # Cuerpo
    tipo_label = {
        "POR OBRA": "por obra",
        "PLAZO FIJO": "a plazo fijo",
        "INDEFINIDO": "indefinido",
    }.get((tipo_contrato or "").upper().strip(), (tipo_contrato or "indefinido").lower())

    fecha_emision_str = f"{fecha_emision.day}-{MESES[fecha_emision.month-1][:3]}-{str(fecha_emision.year)[2:]}"
    fecha_ingreso_str = f"{fecha_ingreso.day:02d} {MESES[fecha_ingreso.month-1]} {fecha_ingreso.year}" if fecha_ingreso else ""

    cuerpo = (
        f"En {ciudad}, a {fecha_emision_str}, {empresa_nombre}, RUT N° {empresa_rut}, "
        f"certifica que {nombre} CI: {rut_empleado}, en su calidad de empleado dependiente, "
        f"desempeñando el cargo de {cargo}, está contratado desde {fecha_ingreso_str}, "
        f"con un contrato {tipo_label}."
    )

    p_cuerpo = doc.add_paragraph()
    p_cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cuerpo.paragraph_format.space_after = Pt(20)
    rc = p_cuerpo.add_run(cuerpo)
    rc.font.size = Pt(11)

    doc.add_paragraph()

    # Cierre
    p_cierre = doc.add_paragraph()
    p_cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_cierre.paragraph_format.space_after = Pt(60)
    rci = p_cierre.add_run("Se extiende el presente certificado para los fines que el solicitante estime necesario.")
    rci.font.size = Pt(11)

    # Punto separador
    p_dot = doc.add_paragraph(".")
    p_dot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dot.paragraph_format.space_after = Pt(60)

    # Firma
    p_firma = doc.add_paragraph("____________________________________")
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_firma.paragraph_format.space_after = Pt(4)
    p_firma.runs[0].font.size = Pt(11)

    p_emp = doc.add_paragraph(empresa_nombre.upper())
    p_emp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_emp.paragraph_format.space_after = Pt(2)
    p_emp.runs[0].font.size = Pt(11)
    p_emp.runs[0].bold = True

    p_rut = doc.add_paragraph(f"C.I. {empresa_rut}")
    p_rut.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rut.runs[0].font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── IRL DS44 ─────────────────────────────────────────────────────────────────

_PROC_F  = ("* Procedimiento Uso herramientas electricas.\n"
            "* procedimiemto instalacion cielo americano\n"
            "* procedimiento trabajo altura\n"
            "* procedimiento instalacion revestimiento metálico")
_PROC_Q  = ("* Procedimiemto instalacion cielo americano\n"
            "* Procedimiemto instalacion revestimiento metálico")
_PROC_B  = ("* Procedimiento Uso, Mantención y Reposición de EPP\n"
            "* RIOHS (última versión)")
_PROC_PS = "RIOHS (última versión)\nprotocolo de RPS"
_PROC_E  = ("* Protocolo TMERT\n"
            "* Procedimiento Uso, Mantención y Reposición de EPP\n"
            "* Procedimiento Uso herramientas electricas.\n"
            "* procedimiemto instalacion cielo americano\n"
            "* procedimiento trabajo altura\n"
            "* Procedimiemto instalacion revestimiento metálico")
_PROC_EM = "* Plan de emergencias"

_IRL_ITEMS = [
    "1. Difusión ley de accidentes y enfermedades profesionales. Ley N°16.744",
    "2. Difusión Política de seguridad y Salud Opcupacional de la empresa.",
    "3. Difusión reglamento interno de Orden Higiene y Seguridad.",
    "4. Correcto uso de elementos de proteccion personal, cuidados y mantenimientos.",
    "5. Difusión matriz de riesgos DS.44",
    "6. Difusión procedimiento como actuar en caso de accidente.",
    "7. Difusión procedimiento de trabajo en altura, metodologia de trabajo correcto.",
    "8. Difusión procedimiento uso de herramientas electricas.",
    "9. Respetar señalizacion al interior de Obra.",
    "10. Desplazamiento por áreas de trabajo señalizadas.",
    "11. Evaluar entorno de trabajo , metodologia de trabajo correcto antes , durante y despues.",
    "12. Prohibición de plataformas de trabajos improvidasas.",
    "13. Autocuidado.",
    "14. Prohibición de ingreso a obra bajo influencia de alcohol y/odrogas.",
    "15. Es obligacion dar aviso inmediato a jefe directo respecto a todo tipo de accidente o lesión que sufra.",
]

_ESPACIO_TRABAJO = (
    "Características generales:\n"
    "Superficies amplias, lisas y a menudo húmedas: se trabaja en superficies de grandes y pequeñas dimensiones que pueden "
    "estar en condiciones despejadas u obstruidas por materiales y estructuras diversas.\n"
    "Ambiente abierto y/o cerrado con ventilación variable: aunque son recintos amplios, pueden tener problemas de recambio "
    "de aire, especialmente si no hay ventilación forzada.\n"
    "Zonas de tránsito de personal y maquinaria.\n"
    "El área debe estar delimitada y señalizada mientras se realizan las obras.\n"
    "Iluminación adecuada (natural o artificial), evitando zonas de penumbra o reflejos en los pisos o techumbres."
)
_COND_AMBIENT = (
    "Idealmente entre 10°C y 30°C, según ficha técnica de los productos con los que se pueda trabajar.\n"
    "Temperaturas extremas pueden afectar el secado y la seguridad del producto aplicado o de los materiales utilizados, "
    "además de afectar al trabajador (hipotermia, estrés térmico).\n"
    "La humedad relativa y la presencia de agua o vapor en el ambiente pueden alterar la adherencia de los productos y/o materiales.\n"
    "Zonas húmedas aumentan el riesgo de resbalones.\n"
    "Ventilación mecánica o natural eficiente es indispensable.\n"
    "En espacios cerrados se deben instalar extractores o equipos de ventilación forzada.\n"
    "Posible exposición a niveles de ruido moderados por maquinaria industrial (evaluar con sonómetro).\n"
    "Si supera los 85 dB(A), debe utilizarse protección auditiva y plan de conservación auditiva.\n"
    "Mínimo 300 lux en áreas de aplicación para garantizar visibilidad de la superficie y detectar imperfecciones.\n"
    "Evitar sombras que puedan dificultar la percepción del piso o de obstáculos."
)
_COND_ORDEN = (
    "Organización de materiales y herramientas: deben almacenarse en zonas designadas, separadas de zonas de tránsito.\n"
    "Señalización y delimitación de áreas: cintas, conos, barreras físicas para restringir el acceso a zonas de trabajo.\n"
    "Evitar cables y mangueras en el suelo sin protección para prevenir caídas.\n"
    "Eliminación de polvo, grasa o restos de productos anteriores.\n"
    "Secado total de la superficie.\n"
    "Limpieza inmediata de derrames químicos según protocolo específico (fichas de datos de seguridad).\n"
    "Se debe realizar una limpieza post-trabajo rigurosa, incluyendo:\n"
    "Lavado de herramientas.\n"
    "Eliminación adecuada de residuos peligrosos (contenedores, restos de productos).\n"
    "Retiro de elementos de señalización cuando el área esté segura."
)
_MAQUINAS = [
    "Aspiradora industrial", "Soldadora", "Esmeril angular de corte",
    "Esmeril angular de desbaste", "Atornillador", "Demoledor manual",
    "Taladro", "Sierra circular", "Sierra sable", "Soplador", "Turbo calefactor",
    "Herramientas manuales (Martillo, Alicate, Pinzas, Llaves variedades, Atornillador, Barreta/Barretilla, Entre otras)",
    "Otras",
]

# (riesgo, daño, medida, proc)
_RIESGOS_FISICOS = [
    ("Caídas al mismo nivel", "Contusión", "Mantener superficies de tránsito en óptimas condiciones, libres de humedad y señalizadas", _PROC_F),
    ("Caídas al mismo nivel", "Fractura(s)", "Transitar solo por zonas demarcadas y habilitadas", _PROC_F),
    ("Caídas al mismo nivel", "Esguince(s)", "Uso de calzado bien ajustado, amarrado y en buen estado", _PROC_F),
    ("Caídas al mismo nivel", "Torceduras", "No utilizar elementos distractores mientras se desplaza", _PROC_F),
    ("Caídas a distinto nivel", "Contusión", "Uso de arnés con doble cabo de vida para trabajos sobre 1,8 metros", _PROC_F),
    ("Caídas a distinto nivel", "Fractura(s)", "Uso de 3 puntos de apoyo al subir y bajar de escaleras o escalas", _PROC_F),
    ("Caídas a distinto nivel", "Esguince(s)", "Mantener superficies de tránsito en óptimas condiciones, libres de humedad y señalizadas", _PROC_F),
    ("Caídas a distinto nivel", "Torceduras", "No detenerse a conversar en escaleras, ni utilizar elementos distractores", _PROC_F),
    ("Caídas de altura", "Contusión", "Uso de 3 puntos de apoyo al subir y bajar de escaleras o escalas", _PROC_F),
    ("Caídas de altura", "Fractura(s)", "Mantener superficies de tránsito en óptimas condiciones, libres de humedad y señalizadas", _PROC_F),
    ("Caídas de altura", "Esguince(s)", "Verificar que la iluminación adecuada para ejecutar las actividades", _PROC_F),
    ("Caídas de altura", "Torceduras", "Contar con examen ocupacional vigente y aprobado para realizar trabajos en altura", _PROC_F),
    ("Caídas de altura", "Muerte", "Utilizar solo plataformas de trabajo certificadas y que posean barandas de protección", _PROC_F),
    ("Atrapamiento", "Herida", "Verificar la ausencia de energías previo a la intervención", _PROC_F),
    ("Atrapamiento", "Corte", "No utilizar ropa holgada, anillos, cadenas o cabello suelto", _PROC_F),
    ("Atrapamiento", "Amputación(es)", "No realizar intervenciones si no se encuentra habilitado para ello", _PROC_F),
    ("Atrapamiento", "Contusión", "Bloqueo de energías. Asegurarse que las maquinarias y herramientas posean todos los sistemas de protección", _PROC_F),
    ("Cortes por objetos / herramientas corto-punzantes", "Corte", "Contar con los EPP necesarios y en buen estado para ejecutar la actividad", _PROC_F),
    ("Cortes por objetos / herramientas corto-punzantes", "Amputación(es)", "Asegurarse que las maquinarias y herramientas posean todos los sistemas de protección", _PROC_F),
    ("Choque contra objetos", "Contusión", "Desplazamiento atento a las condiciones del entorno. No utilizar elementos distractores en el desplazamiento", _PROC_F),
    ("Choque contra objetos", "Herida", "Dispositivos de seguridad de equipos móviles (cámaras, sensores de proximidad, balizas, alarmas de retroceso)", _PROC_F),
    ("Contactos térmicos por calor", "Quemaduras", "Contar con los EPP necesarios y en buen estado para ejecutar la actividad", _PROC_F),
    ("Contactos térmicos por calor", "Herida", "No tocar superficies calientes directamente con las manos", _PROC_F),
    ("Contactos térmicos por frío", "Quemaduras", "Uso de EPP necesarios para la ejecución de la tarea", _PROC_F),
    ("Contactos eléctricos directos baja tensión", "Quemaduras", "No exceder la capacidad de los enchufes", _PROC_F),
    ("Contactos eléctricos directos baja tensión", "Herida", "Desconectar inmediatamente el suministro de energía en caso de cortes circuitos.", _PROC_F),
    ("Contactos eléctricos directos alta tensión", "Asfixia", "Des energizar fuentes de energía previo a ejecución de actividades", _PROC_F),
    ("Contactos eléctricos indirectos baja tensión", "Quemaduras", "Respetar distancias de seguridad entre líneas eléctricas y elementos conductores", _PROC_F),
    ("Contactos eléctricos indirectos alta tensión", "Quemaduras", "Restricción de trabajos próximos a tendido de alta tensión (máximo 15m de próximidad)", _PROC_F),
    ("Contactos eléctricos indirectos alta tensión", "Asfixia", "En caso de arco eléctrico no manipular partes metálicas del equipo que puedan ocasionar descarga eléctrica.", _PROC_F),
    ("Proyección de fragmentos y/o partículas", "Lesión ocular", "No posicionarse en lugares donde exista proyección de partículas.", _PROC_F),
    ("Proyección de fragmentos y/o partículas", "Quemaduras", "Utilización de biombos para evitar la proyección de fragmentos y partículas", _PROC_F),
    ("Proyección de fragmentos y/o partículas", "Dermatitis", "Contar con EPP adecuados para la ejecución de la actividad", _PROC_F),
    ("Atropellos o golpes con vehículos", "Herida", "Dispositivos de seguridad de equipos móviles (cámaras, sensores de proximidad, balizas, sensores de retroceso, luces, alarmas de retroceso)", _PROC_F),
    ("Atropellos o golpes con vehículos", "Contusión", "Respetar normativa de tránsito.", _PROC_F),
    ("Atropellos o golpes con vehículos", "Fractura(s)", "Respetar velocidad máxima permitida. Utilizar solo zonas de tránsito habilitadas", _PROC_F),
    ("Atropellos o golpes con vehículos", "Esguince(s)", "No ingresar a zonas de operación de equipos móviles en operación. Cualquier actividad en estas zonas debe ser coordinada previamente con Jefe del área asegurando la detención del equipo", _PROC_F),
    ("Atropellos o golpes con vehículos", "Muerte", "Está prohibido el acto de interponer nuestro cuerpo o parte de él, entre, frente o bajo una fuente de energía considerablemente mayor a lo que nuestro cuerpo puede soportar (línea de fuego). Ej. trayectoria de equipos móviles, entre otros.", _PROC_F),
    ("Atropellos o golpes con vehículos", "Amputación(es)", "Uso de EPP básicos y chaleco reflectante para visibilidad de personas", _PROC_F),
    ("Choque, colisión o volcamiento", "Herida", "Dispositivos de seguridad de equipos móviles (cámaras, sensores de proximidad, balizas, sensores de retroceso, luces, alarmas de retroceso)", _PROC_F),
    ("Choque, colisión o volcamiento", "Contusión", "Respetar normativa de tránsito.", _PROC_F),
    ("Choque, colisión o volcamiento", "Esguince(s)", "Respetar velocidad máxima permitida. Utilizar solo zonas de tránsito habilitadas", _PROC_F),
    ("Choque, colisión o volcamiento", "Amputación(es)", "Prohibido el uso de elementos distractores mientras conduce", _PROC_F),
    ("Choque, colisión o volcamiento", "Muerte", "No conducir u operar equipos bajo la influencia de drogas, alcohol, estupefacientes u otras sustancias que puedan afectar la concentración o causar somnolencia", _PROC_F),
    ("Choque, colisión o volcamiento", "Fractura(s)", "Uso obligatorio de cinturón de seguridad", _PROC_F),
    ("Exposición a ruido", "Lesión Acústica", "Uso de Protector Auditivo.\nEvaluar periódicamente fuentes generadoras de ruido.\nOperar equipos con puertas y ventanas cerradas.", _PROC_F),
    ("Exposición a Calor", "Colapso físico", "Realizar pausas de trabajo, hidratación abundante, implementación de zonas de sombraje, entre otras)", _PROC_F),
    ("Exposición a Calor", "Dermatitis", "Aplicar al menos cada 2 horas, protector solar FPS 50 en zonas de cara y cuello.", _PROC_F),
    ("Exposición a Frío", "Quemaduras", "Evitar contacto con superficies frías", _PROC_F),
    ("Exposición a Frío", "Colapso físico", "Salas de descanso calefaccionadas", _PROC_F),
    ("Exposición a Frío", "Enfermedad Común (resfrío, gripe)", "Sistema de control de temperatura en equipos móviles", _PROC_F),
    ("Exposición a polvos", "Asma bronquial", "Uso de protección respiratoria (mascarilla con filtros)\nEvaluaciones cualitativas y cuantitativas del agente\nMantención de equipos móviles y sellos herméticos", _PROC_F),
    ("Exposición a polvos", "Lesión ocular", "Uso de lentes de seguridad semisellados y protección facial para labores de limpiezas con aire comprimido.\nOperar y conducir equipos con puertas y ventanas cerradas", _PROC_F),
]

_RIESGOS_QUIMICOS = [
    ("Exposición a sustancias químicas tóxicas", "Conjuntivitis Química", "Contar con todos los EPP adecuados en relación al tipo de sustancia a manipular", _PROC_Q),
    ("Exposición a aerosoles sólidos.", "Asfixia", "Conocer características del producto y formas de manipulación, actuación en caso de contacto, almacenamiento y transporte de acuerdo a lo indicado en HDS", _PROC_Q),
    ("Exposición a aerosoles sólidos.", "Dermatitis", "Almacenamiento de sustancias en estanterías y sobre bandejas para evitar derrames accidentales", _PROC_Q),
    ("Exposición a aerosoles líquidos", "Conjuntivitis Química", "Contar con todos los EPP adecuados en relación al tipo de sustancia a manipular", _PROC_Q),
    ("Exposición a aerosoles líquidos", "Dermatitis", "Conocer características del producto y formas de manipulación, actuación en caso de contacto, almacenamiento y transporte de acuerdo a lo indicado en HDS", _PROC_Q),
    ("Exposición a aerosoles líquidos", "Asfixia", "Almacenamiento de sustancias en estanterías y sobre bandejas para evitar derrames accidentales", _PROC_Q),
    ("Exposición a gases y vapores", "Asfixia", "Contar con todos los EPP adecuados en relación al tipo de sustancia a manipular", _PROC_Q),
    ("Contacto con sustancias cáusticas y/o corrosivas", "Dermatitis", "Contar con todos los EPP adecuados en relación al tipo de sustancia a manipular", _PROC_Q),
    ("Contacto con sustancias cáusticas y/o corrosivas", "Conjuntivitis Química", "Conocer características del producto y formas de manipulación, actuación en caso de contacto, almacenamiento y transporte de acuerdo a lo indicado en HDS", _PROC_Q),
    ("Contacto con sustancias cáusticas y/o corrosivas", "Asfixia", "Almacenamiento de sustancias en estanterías y sobre bandejas para evitar derrames accidentales", _PROC_Q),
]

_RIESGOS_BIOLOGICOS = [
    ("Transmisión por Fluidos Corporales", "Dermatitis", "No compartir utensilios personales.\nDesinfección de zonas de trabajo comunes\nLavado frecuente de manos con agua y jabón\nSanitización de servicios higiénicos y de alimentación", _PROC_B),
    ("Transmisión por Fluidos Corporales", "Enfermedad Común (resfrío, gripe)", "Uso de mascarilla en caso de resfríos y enfermedades que se transmiten por fluidos.", _PROC_B),
    ("Transmisión por Fluidos Corporales", "Alergias", "Limpieza y desinfección de zonas de trabajo comunes", _PROC_B),
    ("Transmisión por Fluidos Corporales", "Enfermedad Común (resfrío, gripe)", "No compartir utensilios personales.\nLavado frecuente de manos con agua y jabón", _PROC_B),
    ("Transmisión por Fluidos Corporales", "Alergias", "Uso de mascarilla en caso de resfríos y enfermedades que se transmiten por fluidos.\nSanitización de servicios higiénicos y de alimentación", _PROC_B),
    ("Transmisión por inhalación dermal, oral y parenteral", "Dermatitis", "Desinfección de zonas de trabajo comunes", _PROC_B),
    ("Contacto con animales y/o insectos", "Herida", "Desinfección, desinsectación y control de plagas en espacios de trabajo", _PROC_B),
    ("Contacto con animales y/o insectos", "Dermatitis", "Establecer programa anual de control de plagas", _PROC_B),
    ("Contacto con animales y/o insectos", "Alergias", "No alimentar animales en lugares de trabajo.\nMantener espacios de trabajo limpios y libres de desechos orgánicos", _PROC_B),
]

_RIESGOS_PSICOSOCIALES = [
    ("Dimensión carga de trabajo (CT)", "Colapso Mental", "Planificación de actividades y organización de tiempo de trabajo.", _PROC_PS),
    ("Dimensión carga de trabajo (CT)", "Neurosis", "Revisión periódica de carga de trabajo y metas.\nEstablecer pausas de trabajo", _PROC_PS),
    ("Dimensión carga de trabajo (CT)", "Colapso físico", "Capacitaciones y conocimientos adecuados a las personas para que puedan desarrollar su trabajo en los tiempos asignados", _PROC_PS),
    ("Dimensión exigencias emocionales (EM)", "Colapso físico", "Tiempos de recuperación y descanso", _PROC_PS),
    ("Dimensión exigencias emocionales (EM)", "Colapso Mental", "Verificar estado de infraestructura física de espacios de trabajo", _PROC_PS),
    ("Dimensión exigencias emocionales (EM)", "Neurosis", "Aplicación y otorgamiento de apoyo organizacional", _PROC_PS),
    ("Dimensión desarrollo profesionales (DP)", "Colapso físico", "Programa de capacitación anual basado en detección de necesidades", _PROC_PS),
    ("Dimensión desarrollo profesionales (DP)", "Colapso Mental", "Programa de promoción interna de trabajadores", _PROC_PS),
    ("Dimensión desarrollo profesionales (DP)", "Neurosis", "Generar instancias de aprendizaje y retroalimentación a los trabajadores", _PROC_PS),
    ("Dimensión reconocimiento y claridad de rol (RC)", "Colapso físico", "Establecer descriptores de cargo y darlos a conocer al personal", _PROC_PS),
    ("Dimensión reconocimiento y claridad de rol (RC)", "Colapso Mental", "Reforzar canales de comunicación interna en la organización", _PROC_PS),
    ("Dimensión reconocimiento y claridad de rol (RC)", "Neurosis", "Establecer instancias de reconocimiento a trabajadores destacados", _PROC_PS),
    ("Dimensión conflicto de rol (CR)", "Colapso físico", "Planificación de trabajo y actividades a realizar", _PROC_PS),
    ("Dimensión conflicto de rol (CR)", "Colapso Mental", "Definir canales de comunicación", _PROC_PS),
    ("Dimensión conflicto de rol (CR)", "Neurosis", "Difundir estructura organizacional", _PROC_PS),
    ("Dimensión calidad del liderazgo (QL)", "Colapso físico", "Formación de lideres en materias de seguridad", _PROC_PS),
    ("Dimensión calidad del liderazgo (QL)", "Colapso Mental", "Generar instancias de retroalimentación bidireccional entre trabajadores y jefaturas", _PROC_PS),
    ("Dimensión calidad del liderazgo (QL)", "Neurosis", "Consensuar normas de respeto a las personas.", _PROC_PS),
    ("Dimensión compañerismo(CM)", "Colapso Mental", "Instancias de recreación grupal (celebraciones de cumpleaños, hitos importantes en faena)", _PROC_PS),
    ("Dimensión compañerismo(CM)", "Neurosis", "Organizar breves encuentros diarios al comenzar la jornada para repartir responsabilidades en forma colectiva.", _PROC_PS),
    ("Dimensión inseguridad en las condiciones de trabajo (IT)", "Colapso físico", "Realizar campañas de buen trato entre las personas.\nFormación a jefaturas.", _PROC_PS),
    ("Dimensión inseguridad en las condiciones de trabajo (IT)", "Colapso Mental", "Consensuar normas de respeto a las personas.", _PROC_PS),
    ("Dimensión inseguridad en las condiciones de trabajo (IT)", "Neurosis", "Promover la organización sindical", _PROC_PS),
]

_RIESGOS_ERGONOMICOS = [
    ("Sobrecarga física debido a la manipulación manual de cargas", "Trastornos músculo-esqueléticos", "Utilización de elementos mecánicos para evitar el MMC.\nRespetar capacidades de carga máximas permitidas legalmente\nCapacitar en técnicas de levantamiento correcto y almacenamiento de materiales", _PROC_E),
    ("Exposición a Vibraciones", "Trastornos músculo-esqueléticos", "Evaluación cualitativa y cuantitativa de los puestos de trabajo\nRealizar Pausas de trabajo y ejercicios compensatorios\nImplementación de Protocolo TMERT", _PROC_E),
    ("Exposición a Vibraciones", "Tendinitis", "Mantención de equipos y herramientas", _PROC_E),
    ("Sobrecarga Postural debido a trabajo de pie", "Trastornos músculo-esqueléticos", "Realizar pausas de trabajo y ejercicios compensatorios\nImplementación de Protocolo TMERT", _PROC_E),
    ("Sobrecarga postural debido a trabajo sentado", "Trastornos músculo-esqueléticos", "Definir zonas de descanso y contar con el personal necesario para la rotación de estos", _PROC_E),
    ("Sobrecarga postural debido a trabajo en cuclillas.", "Trastornos músculo-esqueléticos", "Evaluación cualitativa y cuantitativa de los puestos de trabajo\nRealizar Pausas de trabajo y ejercicios compensatorios\nImplementación de Protocolo TMERT", _PROC_E),
    ("Sobrecarga postural debido a trabajo arrodillado", "Trastornos músculo-esqueléticos", "Evaluación cualitativa y cuantitativa de los puestos de trabajo\nRealizar Pausas de trabajo y ejercicios compensatorios\nImplementación de Protocolo TMERT", _PROC_E),
    ("Sobrecarga Postural debido a Tronco inclinado, en torsión o lateralización", "Trastornos músculo-esqueléticos", "Implementación de Protocolo TMERT", _PROC_E),
    ("Sobrecarga Postural debido a Tronco inclinado, en torsión o lateralización", "Tendinitis", "Realizar pausas de trabajo y ejercicios compensatorios", _PROC_E),
    ("Sobrecarga postural por flexión o extensión de la columna cervical", "Trastornos músculo-esqueléticos", "Definir zonas de descanso y contar con el personal necesario para la rotación de estos", _PROC_E),
    ("Sobrecarga Postural debido a trabajo fuera del alcance funcional", "Trastornos músculo-esqueléticos", "Implementación de Protocolo TMERT", _PROC_E),
    ("Sobrecarga Postural debido a trabajo fuera del alcance funcional", "Tendinitis", "Realizar pausas de trabajo y ejercicios compensatorios.\nDefinir zonas de descanso y contar con el personal necesario para la rotación de estos", _PROC_E),
    ("Sobrecarga postural debido a actividad muscular estática", "Trastornos músculo-esqueléticos", "Implementación de Protocolo TMERT", _PROC_E),
    ("Sobrecarga postural debido a actividad muscular estática", "Tendinitis", "Realizar pausas de trabajo y ejercicios compensatorios.\nDefinir zonas de descanso y contar con el personal necesario para la rotación de estos", _PROC_E),
]

_RIESGOS_EMERGENCIAS = [
    ("Explosiones", "Quemaduras", "Establecer plan de emergencias y evacuación\nDefinir y dar a conocer zonas de seguridad\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas", _PROC_EM),
    ("Explosiones", "Herida", "Establecer plan de emergencias y evacuación\nDefinir y dar a conocer zonas de seguridad\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas", _PROC_EM),
    ("Explosiones", "Lesión Acústica", "Establecer plan de emergencias y evacuación\nDefinir y dar a conocer zonas de seguridad\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas", _PROC_EM),
    ("Incendios", "Quemaduras", "Establecer plan de emergencias y evacuación\nDefinir y dar a conocer zonas de seguridad\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas\nSistemas de extintores tanto en equipos como en espacios de trabajo\nCapacitación uso de extintores", _PROC_EM),
    ("Incendios", "Herida", "Establecer plan de emergencias y evacuación\nDefinir y dar a conocer zonas de seguridad\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas\nSistemas de extintores tanto en equipos como en espacios de trabajo\nCapacitación uso de extintores", _PROC_EM),
    ("Derrame hidrocarburos", "Herida", "Establecer plan de emergencias y evacuación\nBaterías de control de derrames\nPrograma de mantenciones preventivas y correctivas de equipos\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas", _PROC_EM),
    ("Derrame químicos (Diluyentes)", "Herida", "Establecer plan de emergencias y evacuación\nBaterías de control de derrames\nCapacitación en Procedimientos de almacenamiento de materiales\nBandejas de contención de residuos para almacenamiento de materiales\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas", _PROC_EM),
    ("Manejo inadecuado de residuos", "Herida", "Establecer plan de emergencias y evacuación\nBaterías de control de derrames\nProcedimiento de manejos de residuos estableciendo identificación, clasificación, segregación y eliminación\nPlan de simulacros anual\nRealizar simulacros y establecer medidas correctivas y preventivas", _PROC_EM),
    ("Transporte inadecuado de sustancias peligrosas", "Herida", "Establecer plan de emergencias y evacuación\nBaterías de control de derrames\nHabilitación y certificación de vehículos para transporte de SUSPEL\nCapacitación en HDS de SUSPEL", _PROC_EM),
    ("Emisiones de aire", "Herida", "Revisión técnica de equipos\nCertificación de equipos móviles\nPrograma de mantención de equipos móviles", _PROC_EM),
]


def _irl_risk_table(doc, categoria: str, filas: list):
    """Agrega una tabla de riesgos al documento IRL."""
    hdr_bg = "D9D9D9"
    tbl = doc.add_table(rows=1 + len(filas), cols=4)
    tbl.style = "Table Grid"
    tbl.autofit = False
    from docx.shared import Cm
    widths = [Cm(3.8), Cm(2.8), Cm(6.0), Cm(5.4)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]

    # Header row
    hdr = tbl.rows[0]
    headers = [
        categoria,
        "Daño Potencial",
        "Medidas Preventivas (Eliminación;\nSustitución; Controles Ing. Y adm.; EPP",
        "Capacitación/Procedimiento",
    ]
    for i, (cell, txt) in enumerate(zip(hdr.cells, headers)):
        _set_cell_bg(cell, hdr_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(7)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, (riesgo, dano, medida, proc) in enumerate(filas):
        row = tbl.rows[r_idx + 1]
        for i, (cell, txt) in enumerate(zip(row.cells, [riesgo, dano, medida, proc])):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(txt)
            run.font.size = Pt(7)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i <= 1 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()


def generar_irl_docx(
    nombre_trabajador: str,
    rut_trabajador: str,
    cargo: str,
    obra_nombre: str,
    obra_direccion: str,
    fecha,
    hora_inicio: str,
    hora_termino: str,
    relator_nombre: str,
    relator_cargo: str,
    empresa=None,
) -> bytes:
    doc = Document()

    # Márgenes
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    sec = doc.sections[0]
    sec.page_width  = Cm(21.59)
    sec.page_height = Cm(27.94)
    for attr, val in [("top","1.5cm"),("bottom","1.5cm"),("left","2cm"),("right","1.5cm")]:
        setattr(sec, f"{attr}_margin", Cm(float(val[:-2])))

    # ── Header ──
    hdr_tbl = doc.add_table(rows=1, cols=3)
    hdr_tbl.style = "Table Grid"
    hdr_tbl.autofit = False
    widths_h = [Cm(3.5), Cm(11.0), Cm(3.5)]
    for i, cell in enumerate(hdr_tbl.rows[0].cells):
        cell.width = widths_h[i]

    # Logo
    logo_cell = hdr_tbl.rows[0].cells[0]
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    logo_url = getattr(empresa, "logo_url", None) if empresa else None
    _add_logo(logo_cell, logo_url)

    # Título centro
    title_cell = hdr_tbl.rows[0].cells[1]
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_lines = ["INFORMACIÓN DE RIESGOS LABORALES", "INDUCCION TRABAJADOR NUEVO", "Artículo N°15 - Decreto Supremo N°44"]
    # reusar el párrafo vacío inicial, luego agregar el resto
    for idx, line in enumerate(title_lines):
        p = title_cell.paragraphs[0] if idx == 0 else title_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(10)

    # Versión
    ver_cell = hdr_tbl.rows[0].cells[2]
    ver_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for idx, line in enumerate(["Versión 0", "Agosto 2025"]):
        p = ver_cell.paragraphs[0] if idx == 0 else ver_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(8)

    doc.add_paragraph()

    # ── Texto legal ──
    legal = doc.add_paragraph()
    legal.paragraph_format.space_after = Pt(6)
    run = legal.add_run(
        '"La entidad empleadora deberá garantizar que cada persona trabajadora, previo al inicio de las labores, '
        'reciba de forma oportuna y adecuada información acerca de los riesgos que entrañan sus labores, de las '
        'medidas preventivas y los métodos o procedimientos de trabajo correctos, los riesgos son los inherentes '
        'a la actividad de cada empresa".'
    )
    run.font.size = Pt(8)
    legal.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ── Datos Generales ──
    def _sec_title(txt):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.style = "Table Grid"
        cell = tbl.rows[0].cells[0]
        _set_cell_bg(cell, "D9D9D9")
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(txt)
        run.bold = True
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return tbl

    _sec_title("DATOS GENERALES")

    dg = doc.add_table(rows=4, cols=4)
    dg.style = "Table Grid"
    dg.autofit = False
    for row in dg.rows:
        for cell in row.cells:
            cell.width = Cm(4.5)

    def _dg(row, col, label, value, bold_label=True):
        cell = dg.rows[row].cells[col]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.clear()
        r = p.add_run(label)
        r.bold = bold_label
        r.font.size = Pt(8)
        if value:
            r2 = p.add_run(f"  {value}")
            r2.font.size = Pt(8)

    # Row 0: obra span (merge cols 0-1) | dirección span (2-3)
    dg.rows[0].cells[0].merge(dg.rows[0].cells[1])
    dg.rows[0].cells[2].merge(dg.rows[0].cells[3])
    _dg(0, 0, f"OBRA: {obra_nombre}", "")
    _dg(0, 2, f"DIRECCION: {obra_direccion}", "")

    _dg(1, 0, "Nombre trabajador", "")
    dg.rows[1].cells[0].merge(dg.rows[1].cells[1])
    c = dg.rows[1].cells[0]
    c.paragraphs[0].clear()
    r = c.paragraphs[0].add_run("Nombre trabajador")
    r.bold = True; r.font.size = Pt(8)
    r2 = c.paragraphs[0].add_run(f"  {nombre_trabajador}")
    r2.bold = True; r2.font.size = Pt(8)

    _dg(1, 2, "Rut", f"  {rut_trabajador}")

    fecha_str = fecha.strftime("%d-%m-%Y") if hasattr(fecha, "strftime") else str(fecha)
    _dg(2, 0, "Fecha", fecha_str)
    _dg(2, 2, "Cargo/Especialidad", cargo)

    _dg(3, 0, "Hora inicio", hora_inicio)
    _dg(3, 2, "Hora término", hora_termino)

    doc.add_paragraph()

    # ── De la información de los riesgos laborales ──
    _sec_title("DE LA INFORMACION DE LOS RIESGOS LABORALES")
    items_tbl = doc.add_table(rows=len(_IRL_ITEMS), cols=1)
    items_tbl.style = "Table Grid"
    for i, item in enumerate(_IRL_ITEMS):
        cell = items_tbl.rows[i].cells[0]
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(item)
        run.font.size = Pt(8)

    doc.add_paragraph()

    # ── Características del lugar de trabajo ──
    _sec_title("CARACTERISTICAS DEL LUGAR DE TRABAJO")

    char_tbl = doc.add_table(rows=5, cols=2)
    char_tbl.style = "Table Grid"
    char_tbl.autofit = False
    char_tbl.rows[0].cells[0].width = Cm(4)
    char_tbl.rows[0].cells[1].width = Cm(14)

    char_data = [
        ("Equipo", "No Aplica\t\tMarca-Modelo\t\tNo Aplica"),
        ("Espacio de Trabajo", _ESPACIO_TRABAJO),
        ("Condiciones Ambientales del Puesto de Trabajo", _COND_AMBIENT),
        ("Condiciones de Orden y Aseo exigidas en el Lugar de Trabajo", _COND_ORDEN),
        ("Máquinas o herramientas que se deben emplear", "\n".join(f"☐  {m}" for m in _MAQUINAS)),
    ]
    for i, (label, content) in enumerate(char_data):
        lc = char_tbl.rows[i].cells[0]
        lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = lc.paragraphs[0]
        p.clear()
        run = p.add_run(label)
        run.bold = True; run.font.size = Pt(7.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cc = char_tbl.rows[i].cells[1]
        cc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p2 = cc.paragraphs[0]
        p2.clear()
        run2 = p2.add_run(content)
        run2.font.size = Pt(7.5)

    doc.add_paragraph()

    # ── Riesgos Específicos ──
    rt = doc.add_table(rows=1, cols=1)
    rt.style = "Table Grid"
    _set_cell_bg(rt.rows[0].cells[0], "D9D9D9")
    p = rt.rows[0].cells[0].paragraphs[0]
    p.clear()
    run = p.add_run("RIESGOS ESPECÍFICOS , DAÑO POTENCIAL , MEDIDAS DE CONTROL Y CAPACITACIÓN")
    run.bold = True; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    _irl_risk_table(doc, "Riesgos Físicos", _RIESGOS_FISICOS)
    _irl_risk_table(doc, "Riesgos Químicos", _RIESGOS_QUIMICOS)
    _irl_risk_table(doc, "Riesgos Biológicos", _RIESGOS_BIOLOGICOS)
    _irl_risk_table(doc, "Riesgos Psicosociales", _RIESGOS_PSICOSOCIALES)
    _irl_risk_table(doc, "Riesgos Ergonómicos", _RIESGOS_ERGONOMICOS)
    _irl_risk_table(doc, "Riesgos por Emergencias", _RIESGOS_EMERGENCIAS)

    # ── Consentimiento del trabajador ──
    ct = doc.add_table(rows=2, cols=3)
    ct.style = "Table Grid"
    ct.autofit = False
    _set_cell_bg(ct.rows[0].cells[0], "D9D9D9")
    ct.rows[0].cells[0].merge(ct.rows[0].cells[1])
    ct.rows[0].cells[0].merge(ct.rows[0].cells[2])
    p = ct.rows[0].cells[0].paragraphs[0]
    p.clear()
    run = p.add_run("CONSENTIMIENTO DEL TRABAJADOR")
    run.bold = True; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 1: nombre | rut | firma
    c0 = ct.rows[1].cells[0]
    p = c0.paragraphs[0]; p.clear()
    r = p.add_run("Nombre del Trabajador : "); r.bold = True; r.font.size = Pt(8)
    r2 = p.add_run(nombre_trabajador); r2.font.size = Pt(8)

    c1 = ct.rows[1].cells[1]
    p = c1.paragraphs[0]; p.clear()
    r = p.add_run("Rut: "); r.bold = True; r.font.size = Pt(8)
    r2 = p.add_run(rut_trabajador); r2.font.size = Pt(8)

    c2 = ct.rows[1].cells[2]
    p = c2.paragraphs[0]; p.clear()
    r = p.add_run("Firma del Trabajador"); r.bold = True; r.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # blank lines for signature
    for _ in range(3):
        c2.add_paragraph()

    doc.add_paragraph()

    # ── Datos Relator ──
    dr = doc.add_table(rows=2, cols=3)
    dr.style = "Table Grid"
    _set_cell_bg(dr.rows[0].cells[0], "D9D9D9")
    dr.rows[0].cells[0].merge(dr.rows[0].cells[1])
    dr.rows[0].cells[0].merge(dr.rows[0].cells[2])
    p = dr.rows[0].cells[0].paragraphs[0]; p.clear()
    run = p.add_run("DATOS RELATOR"); run.bold = True; run.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c0 = dr.rows[1].cells[0]
    p = c0.paragraphs[0]; p.clear()
    r = p.add_run("Nombre Relator : "); r.bold = True; r.font.size = Pt(8)
    r2 = p.add_run(relator_nombre); r2.font.size = Pt(8)

    c1 = dr.rows[1].cells[1]
    p = c1.paragraphs[0]; p.clear()
    r = p.add_run("Cargo : "); r.bold = True; r.font.size = Pt(8)
    r2 = p.add_run(relator_cargo); r2.font.size = Pt(8)

    c2 = dr.rows[1].cells[2]
    p = c2.paragraphs[0]; p.clear()
    r = p.add_run("Firma Relator"); r.bold = True; r.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(3):
        c2.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
