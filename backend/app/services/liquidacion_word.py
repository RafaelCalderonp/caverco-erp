import io
from decimal import Decimal

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MESES = ["enero","febrero","marzo","abril","mayo","junio",
         "julio","agosto","septiembre","octubre","noviembre","diciembre"]

# ── helpers ──────────────────────────────────────────────────────────────────

def _clp(v) -> str:
    if v is None or v == 0:
        return "$ -"
    return f"$ {int(v):,}".replace(",", ".")

def _numero_letras(n: int) -> str:
    if n <= 0:
        return "Cero"
    unidades = ["","Un","Dos","Tres","Cuatro","Cinco","Seis","Siete","Ocho","Nueve",
                "Diez","Once","Doce","Trece","Catorce","Quince","Dieciséis","Diecisiete",
                "Dieciocho","Diecinueve"]
    decenas  = ["","","Veinte","Treinta","Cuarenta","Cincuenta","Sesenta","Setenta","Ochenta","Noventa"]
    centenas = ["","Cien","Doscientos","Trescientos","Cuatrocientos","Quinientos",
                "Seiscientos","Setecientos","Ochocientos","Novecientos"]

    def _grupo(x):
        if x == 0: return ""
        if x == 100: return "Cien"
        c, resto = divmod(x, 100)
        partes = []
        if c: partes.append(centenas[c])
        if 0 < resto < 20:
            partes.append(unidades[resto])
        elif resto >= 20:
            d, u = divmod(resto, 10)
            partes.append(decenas[d] + (" y " + unidades[u] if u else ""))
        return " ".join(p for p in partes if p)

    partes = []
    millones, resto = divmod(n, 1_000_000)
    miles,    unid  = divmod(resto, 1_000)
    if millones == 1:  partes.append("Un Millón")
    elif millones > 1: partes.append(_grupo(millones) + " Millones")
    if miles == 1:     partes.append("Mil")
    elif miles > 1:    partes.append(_grupo(miles) + " Mil")
    if unid:           partes.append(_grupo(unid))
    return " ".join(p for p in partes if p)

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_row_height(row, cm: float):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(cm * 567)))   # 1 cm ≈ 567 twips
    trHeight.set(qn("w:hRule"), "exact")
    old = trPr.find(qn("w:trHeight"))
    if old is not None: trPr.remove(old)
    trPr.append(trHeight)

def _set_borders(tbl, color="CCCCCC", sz="4"):
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tblBorders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None: tblPr.remove(old)
    tblPr.append(tblBorders)

def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    old = tcPr.find(qn("w:tcMar"))
    if old is not None: tcPr.remove(old)
    tcPr.append(tcMar)

def _para(cell, text, bold=False, size=9, align=None, color=None):
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return p

def _remove_paragraph_spacing(doc):
    """Elimina el espaciado automático entre párrafos del documento."""
    for style in doc.styles:
        if hasattr(style, 'paragraph_format'):
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after  = Pt(0)

def _spacer(doc, pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)
    run = p.add_run(" ")
    run.font.size = Pt(pts)


# ── generador principal ───────────────────────────────────────────────────────

def generar_liquidacion_docx(empresa, empleado, liquidacion,
                              afp_nombre: str, isapre_nombre: str,
                              cargo_nombre: str, centro_costo_codigo: str,
                              fecha_ingreso, logo_bytes) -> bytes:
    doc = Document()
    _remove_paragraph_spacing(doc)

    # Márgenes carta (21.59 × 27.94 cm)
    for sec in doc.sections:
        sec.page_width  = Cm(21.59)
        sec.page_height = Cm(27.94)
        sec.top_margin    = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    HEADER  = "475569"
    AZUL    = "1E3A5F"
    GRIS    = "F2F4F6"
    TOTAL_W = Cm(17.59)   # ancho disponible (21.59 - 2*2.0)

    clp = _clp
    periodo = liquidacion.periodo       # YYYY-MM
    anio, mes_num = periodo.split("-")
    mes_label = MESES[int(mes_num) - 1].upper()

    empresa_nombre = (empresa.razon_social if empresa else None) or ""

    # ── 1. Encabezado: [empresa | logo] ──────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=2)
    hdr.autofit = False
    hdr.columns[0].width = Cm(11)
    hdr.columns[1].width = Cm(6.59)

    # Nombre empresa izquierda
    p_emp = hdr.rows[0].cells[0].paragraphs[0]
    p_emp.paragraph_format.space_before = Pt(0)
    p_emp.paragraph_format.space_after  = Pt(0)
    r_emp = p_emp.add_run(empresa_nombre)
    r_emp.bold = True
    r_emp.font.size = Pt(11)
    hdr.rows[0].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Logo derecha
    logo_cell = hdr.rows[0].cells[1]
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after  = Pt(0)
    if logo_bytes:
        try:
            p_logo.add_run().add_picture(io.BytesIO(logo_bytes), height=Cm(1.2))
        except Exception:
            pass  # logo inválido — celda queda vacía

    _spacer(doc, 4)

    # ── 2. Título centrado ────────────────────────────────────────────────────
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.space_before = Pt(0)
    p_tit.paragraph_format.space_after  = Pt(0)
    r_tit = p_tit.add_run("LIQUIDACIÓN DE REMUNERACIONES")
    r_tit.bold = True
    r_tit.font.size = Pt(13)
    r_tit.font.color.rgb = RGBColor.from_string(AZUL)

    _spacer(doc, 3)

    # ── 3. Período ────────────────────────────────────────────────────────────
    p_per = doc.add_paragraph()
    p_per.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_per.paragraph_format.space_before = Pt(0)
    p_per.paragraph_format.space_after  = Pt(0)
    rp = p_per.add_run(f"PERÍODO DE REMUNERACIONES:  {mes_label} {anio}")
    rp.bold = True
    rp.font.size = Pt(9)

    _spacer(doc, 4)

    # ── 4. Datos del trabajador ───────────────────────────────────────────────
    nombre_completo = f"{empleado.nombres} {empleado.apellido_paterno} {empleado.apellido_materno or ''}".strip()
    tasa_afp_label  = afp_nombre.upper() if afp_nombre else "—"
    isapre_label    = isapre_nombre.upper() if isapre_nombre else "—"
    uf_val          = float(liquidacion.valor_uf or 0)
    uf_label        = f"( {uf_val:.2f} UF )" if uf_val else ""

    datos = [
        (None, "DATOS TRABAJADOR"),          # fila título
        ("NOMBRE:",              nombre_completo),
        ("RUT:",                 empleado.rut),
        ("CARGO:",               cargo_nombre),
        ("INSTITUCIÓN PREVISIONAL:", tasa_afp_label),
        ("INSTITUCIÓN SALUD:",   f"{isapre_label} {uf_label}".strip()),
        ("CENTRO DE COSTO:",     centro_costo_codigo or "—"),
        ("DÍAS TRABAJADOS:",     str(liquidacion.dias_trabajados or 30)),
        ("CARGAS:",              str(empleado.n_cargas or 0)),
        ("FECHA DE INGRESO:",    fecha_ingreso.strftime("%d-%m-%Y") if fecha_ingreso else "—"),
    ]

    ROW_H = 0.52   # cm — altura compacta por fila
    dt = doc.add_table(rows=len(datos), cols=2)
    _set_borders(dt, "CCCCCC", "4")
    dt.autofit = False
    dt.columns[0].width = Cm(6)
    dt.columns[1].width = Cm(11.59)

    for i, (label, valor) in enumerate(datos):
        row = dt.rows[i]
        _set_row_height(row, ROW_H)
        if label is None:
            # Fila título de sección
            merged = row.cells[0].merge(row.cells[1])
            _set_cell_bg(merged, HEADER)
            _set_cell_margins(merged, top=30, bottom=30)
            _para(merged, valor, bold=True, size=9, color="FFFFFF")
        else:
            _set_cell_bg(row.cells[0], GRIS)
            _set_cell_bg(row.cells[1], "FFFFFF")
            _set_cell_margins(row.cells[0], top=30, bottom=30)
            _set_cell_margins(row.cells[1], top=30, bottom=30)
            _para(row.cells[0], label, bold=True, size=9)
            _para(row.cells[1], valor, size=9)

    _spacer(doc, 4)

    # ── 5. Base imponible / tributable ────────────────────────────────────────
    bi = doc.add_table(rows=1, cols=4)
    _set_borders(bi, "CCCCCC", "4")
    bi.autofit = False
    for i, w in enumerate([Cm(3.5), Cm(5.3), Cm(3.5), Cm(5.29)]):
        bi.columns[i].width = w
    _set_row_height(bi.rows[0], ROW_H)
    _set_cell_bg(bi.rows[0].cells[0], GRIS)
    _set_cell_bg(bi.rows[0].cells[2], GRIS)
    for c in bi.rows[0].cells:
        _set_cell_margins(c, top=30, bottom=30)
    _para(bi.rows[0].cells[0], "BASE IMPONIBLE",  bold=True, size=8)
    _para(bi.rows[0].cells[1], clp(liquidacion.total_imponible), size=9,
          align=WD_ALIGN_PARAGRAPH.RIGHT)
    _para(bi.rows[0].cells[2], "BASE TRIBUTABLE", bold=True, size=8)
    _para(bi.rows[0].cells[3], clp(liquidacion.base_tributaria), size=9,
          align=WD_ALIGN_PARAGRAPH.RIGHT)

    _spacer(doc, 4)

    # ── 6. Título tabla principal ─────────────────────────────────────────────
    p_det = doc.add_paragraph()
    p_det.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_det.paragraph_format.space_before = Pt(0)
    p_det.paragraph_format.space_after  = Pt(0)
    rd = p_det.add_run("DETALLE CONCEPTOS HABER Y DESCUENTO")
    rd.bold = True
    rd.font.size = Pt(9)
    rd.font.color.rgb = RGBColor.from_string(AZUL)

    # ── 7. Tabla principal 3 columnas ─────────────────────────────────────────
    COL_W = Cm(17.59 / 3)
    main = doc.add_table(rows=1, cols=3)
    _set_borders(main, "CCCCCC", "4")
    main.autofit = False
    for col in main.columns:
        col.width = COL_W

    # Encabezados
    _set_row_height(main.rows[0], ROW_H)
    for cell, txt in zip(main.rows[0].cells,
                         ["HABERES IMPONIBLES", "HABERES NO IMPONIBLES", "DESCUENTOS"]):
        _set_cell_bg(cell, HEADER)
        _set_cell_margins(cell, top=30, bottom=30)
        _para(cell, txt, bold=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, color="FFFFFF")

    liq = liquidacion

    def _fila3(l0, v0, l1, v1, l2, v2):
        row = main.add_row()
        _set_row_height(row, ROW_H)
        for cell, lbl, val in zip(row.cells, [l0, l1, l2], [v0, v1, v2]):
            _set_cell_margins(cell, top=25, bottom=25)
            p = cell.paragraphs[0]
            p.clear()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            if lbl:
                r1 = p.add_run(lbl + "  ")
                r1.font.size = Pt(8)
            if val:
                r2 = p.add_run(val)
                r2.font.size = Pt(8)
                r2.bold = True
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _fila3("SUELDO BASE",        clp(liq.sueldo_base),
           "MOVILIZACIÓN",       clp(liq.movilizacion),
           "COTIZ. AFP",         clp(liq.descuento_afp))
    _fila3("GRATIFICACIÓN LEGAL",clp(liq.gratificacion),
           "COLACIÓN",           clp(liq.colacion),
           "INST. SALUD",        clp(liq.descuento_salud))
    _fila3("HH.EE 50%"  if liq.horas_extra_50  else "",
           clp(liq.horas_extra_50)  if liq.horas_extra_50  else "",
           "VIÁTICO",            clp(liq.viaticos),
           "ADIC. SALUD",        clp(liq.adicional_salud) if liq.adicional_salud else "$ -")
    _fila3("HH.EE 100%" if liq.horas_extra_100 else "",
           clp(liq.horas_extra_100) if liq.horas_extra_100 else "",
           "FAMILIAR",           clp(liq.asig_familiar) if liq.asig_familiar else "$ -",
           "SEG. CESANTÍA",      clp(liq.afc_trabajador))
    _fila3("AGUINALDO"  if liq.aguinaldo        else "",
           clp(liq.aguinaldo)        if liq.aguinaldo        else "",
           "",                   "",
           "IMPUESTO ÚNICO",     clp(liq.impuesto_unico))

    # Fila totales
    row_tot = main.add_row()
    _set_row_height(row_tot, ROW_H)
    no_imp = (liq.total_haberes or 0) - (liq.total_imponible or 0)
    for c, (lbl, val) in enumerate([
        ("TOTAL IMPONIBLES",    clp(liq.total_imponible)),
        ("TOTAL NO IMPONIBLES", clp(no_imp)),
        ("TOTAL DESCUENTOS",    clp(liq.total_desc_legales)),
    ]):
        _set_cell_bg(row_tot.cells[c], GRIS)
        _set_cell_margins(row_tot.cells[c], top=30, bottom=30)
        p = row_tot.cells[c].paragraphs[0]
        p.clear()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r1 = p.add_run(lbl + "  ")
        r1.bold = True; r1.font.size = Pt(8)
        r2 = p.add_run(val)
        r2.bold = True; r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor.from_string(AZUL)
        row_tot.cells[c].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _spacer(doc, 4)

    # ── 8. Alcance líquido ────────────────────────────────────────────────────
    total_pagar = max(0, int(liq.liquido_a_pagar or 0) - int(liq.anticipo or 0))
    liq_tbl = doc.add_table(rows=3, cols=2)
    _set_borders(liq_tbl, "CCCCCC", "4")
    liq_tbl.autofit = False
    liq_tbl.columns[0].width = Cm(10)
    liq_tbl.columns[1].width = Cm(7.59)

    for i, (lbl, val) in enumerate([
        ("ALCANCE LÍQUIDO", clp(liq.liquido_a_pagar)),
        ("ANTICIPO",        clp(liq.anticipo)),
        ("TOTAL A PAGAR",   clp(total_pagar)),
    ]):
        row = liq_tbl.rows[i]
        _set_row_height(row, ROW_H)
        is_total = i == 2
        bg = HEADER if is_total else GRIS
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], HEADER if is_total else "FFFFFF")
        for c in row.cells:
            _set_cell_margins(c, top=30, bottom=30)
        color = "FFFFFF" if is_total else "3B3B3B"
        _para(row.cells[0], lbl, bold=True, size=9, color=color)
        _para(row.cells[1], val, bold=True, size=10,
              align=WD_ALIGN_PARAGRAPH.RIGHT, color=color)

    _spacer(doc, 3)

    # ── 9. SON ────────────────────────────────────────────────────────────────
    p_son = doc.add_paragraph()
    p_son.paragraph_format.space_before = Pt(0)
    p_son.paragraph_format.space_after  = Pt(0)
    rs = p_son.add_run(f"SON:  {_numero_letras(total_pagar)} pesos.")
    rs.font.size = Pt(8)
    rs.italic = True

    _spacer(doc, 6)

    # ── 10. Firma ─────────────────────────────────────────────────────────────
    firma = doc.add_table(rows=2, cols=2)
    firma.autofit = False
    firma.columns[0].width = Cm(9.5)
    firma.columns[1].width = Cm(8.09)

    p_cert = firma.rows[0].cells[0].paragraphs[0]
    p_cert.paragraph_format.space_before = Pt(0)
    p_cert.paragraph_format.space_after  = Pt(0)
    r_cert = p_cert.add_run(
        "Certifico que he recibido a entera satisfacción los valores contenidos en la "
        "presente liquidación de sueldo, por lo cual no tengo ningún cargo o reclamo "
        "posterior que efectuar a mi empleador."
    )
    r_cert.font.size = Pt(7.5)

    p_rc = firma.rows[0].cells[1].paragraphs[0]
    p_rc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rc.paragraph_format.space_before = Pt(0)
    p_rc.paragraph_format.space_after  = Pt(0)
    r_rc = p_rc.add_run("RECIBI CONFORME")
    r_rc.bold = True
    r_rc.font.size = Pt(9)

    for c in range(2):
        p = firma.rows[1].cells[c].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.add_run("_" * 38).font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
